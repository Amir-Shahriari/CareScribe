"""Filling a template: header values land inline or as a new paragraph
depending on style, field values land in their mapped cell, and the
document's structure (table/row counts, section headers, signature row)
survives untouched.
"""

import io

import docx
import pytest

from carescribe.core import clinical_forms


def test_fill_template_writes_header_and_field_values():
    spec = clinical_forms.get_form_spec("client_session_notes")
    header_values = {
        "date": "12/08/2026", "practitioner": "A. Nguyen",
        "client_name": "J. Smith", "client_dob": "01/01/1990",
        "session_number": "4", "item_code": "80010",
        "reason_for_referral": "Ongoing anxiety management.",
    }
    field_values = {f.key: f"Generated text for {f.key}." for f in spec.fields}

    output = clinical_forms.fill_template(spec, field_values, header_values)
    doc = docx.Document(io.BytesIO(output))

    row0 = clinical_forms._dedupe_row(doc.tables[0].rows[0])
    assert row0[0].text == "Date: 12/08/2026"
    assert row0[1].text == "Practitioner: A. Nguyen"

    reason_cell = clinical_forms._dedupe_row(doc.tables[0].rows[4])[0]
    assert "Reason for referral:" in reason_cell.paragraphs[0].text
    assert "Ongoing anxiety management." in reason_cell.text

    sample_field = spec.fields[0]
    value_cell = clinical_forms._dedupe_row(
        doc.tables[sample_field.table_index].rows[sample_field.value_row_index]
    )[sample_field.value_col_index]
    assert f"Generated text for {sample_field.key}." in value_cell.text


def test_fill_template_preserves_structure():
    spec = clinical_forms.get_form_spec("client_session_notes")
    field_values = {f.key: "x" for f in spec.fields}
    output = clinical_forms.fill_template(spec, field_values, {})
    original = docx.Document(spec.asset_path)
    filled = docx.Document(io.BytesIO(output))

    assert len(filled.tables) == len(original.tables)
    assert len(filled.tables[0].rows) == len(original.tables[0].rows)
    # Section header untouched.
    assert filled.tables[0].rows[6].cells[0].text == "SESSION SUMMARY"
    # Signature row untouched (still blank).
    sig_row = clinical_forms._dedupe_row(filled.tables[0].rows[18])
    assert sig_row[0].text == "Signature:"


def test_fill_template_defaults_missing_field_to_not_documented():
    spec = clinical_forms.get_form_spec("client_session_notes")
    output = clinical_forms.fill_template(spec, {}, {})
    doc = docx.Document(io.BytesIO(output))
    sample_field = spec.fields[0]
    value_cell = clinical_forms._dedupe_row(
        doc.tables[sample_field.table_index].rows[sample_field.value_row_index]
    )[sample_field.value_col_index]
    assert value_cell.text.strip() == "Not documented"


def test_fill_template_defaults_whitespace_only_field_to_not_documented():
    # A field value of only \r/\t (e.g. surviving a generation glitch) must
    # not be written verbatim — it should default the same as a missing one.
    spec = clinical_forms.get_form_spec("client_session_notes")
    field_values = {f.key: "\r\n\r" for f in spec.fields}
    output = clinical_forms.fill_template(spec, field_values, {})
    doc = docx.Document(io.BytesIO(output))
    sample_field = spec.fields[0]
    value_cell = clinical_forms._dedupe_row(
        doc.tables[sample_field.table_index].rows[sample_field.value_row_index]
    )[sample_field.value_col_index]
    assert value_cell.text.strip() == "Not documented"


@pytest.mark.parametrize(
    "form_id",
    ["client_session_notes", "client_treatment_review", "biopsychosocial_assessment"],
)
def test_fill_template_writes_each_header_field_to_its_own_cell(form_id):
    # Regression guard for a future template edit that swaps which physical
    # cell a header label sits in: fill every header field with a value
    # unique to that field, then confirm each sentinel landed in exactly the
    # cell its own FormSpec entry names — not just "somewhere in the
    # document" and not just "the count is right". If two header entries
    # were ever mis-specified to point at the same cell, both sentinels
    # would land there and the cross-check below would catch it.
    spec = clinical_forms.get_form_spec(form_id)
    header_values = {header.key: f"SENTINEL_{header.key}" for header in spec.header_fields}
    output = clinical_forms.fill_template(spec, {}, header_values)
    doc = docx.Document(io.BytesIO(output))

    cell_text_by_key = {
        header.key: clinical_forms._dedupe_row(
            doc.tables[header.table_index].rows[header.row_index]
        )[header.col_index].text
        for header in spec.header_fields
    }

    for header in spec.header_fields:
        own_text = cell_text_by_key[header.key]
        assert f"SENTINEL_{header.key}" in own_text
        for other in spec.header_fields:
            if other.key != header.key:
                assert f"SENTINEL_{other.key}" not in own_text


def test_fill_cell_defaults_whitespace_only_text_to_not_documented():
    spec = clinical_forms.get_form_spec("client_session_notes")
    doc = docx.Document(spec.asset_path)
    sample_field = spec.fields[0]
    cell = clinical_forms._dedupe_row(
        doc.tables[sample_field.table_index].rows[sample_field.value_row_index]
    )[sample_field.value_col_index]
    clinical_forms._fill_cell(cell, "\t")
    assert cell.text.strip() == "Not documented"


def test_fill_template_raises_clinical_form_error_on_header_shape_drift():
    # A header whose row index no longer exists in the template (e.g. after
    # someone edits the .docx and rows shift) must fail loudly with a
    # ClinicalFormError identifying WHERE it broke, not a bare IndexError —
    # and must never echo the value being written into the error message.
    spec = clinical_forms.get_form_spec("client_session_notes")
    bad_header = clinical_forms.HeaderField(
        key="bogus", label="Bogus", table_index=0,
        row_index=9999, col_index=0, style="inline",
    )
    bad_spec = clinical_forms.FormSpec(
        form_id=spec.form_id, title=spec.title, asset_path=spec.asset_path,
        header_fields=[bad_header], fields=[],
    )
    with pytest.raises(clinical_forms.ClinicalFormError) as excinfo:
        clinical_forms.fill_template(bad_spec, {}, {"bogus": "secret header value"})
    message = str(excinfo.value)
    assert spec.form_id in message
    assert "9999" in message
    assert "secret header value" not in message


def test_fill_template_raises_clinical_form_error_on_field_shape_drift():
    spec = clinical_forms.get_form_spec("client_session_notes")
    bad_field = clinical_forms.FormField(
        key="bogus", label="Bogus", table_index=0,
        value_row_index=9999, value_col_index=0, append_after_label=False,
    )
    bad_spec = clinical_forms.FormSpec(
        form_id=spec.form_id, title=spec.title, asset_path=spec.asset_path,
        header_fields=[], fields=[bad_field],
    )
    with pytest.raises(clinical_forms.ClinicalFormError) as excinfo:
        clinical_forms.fill_template(bad_spec, {"bogus": "secret clinical content"}, {})
    message = str(excinfo.value)
    assert spec.form_id in message
    assert "9999" in message
    assert "secret clinical content" not in message
