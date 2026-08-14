"""Generic table-row classification: which rows are fields, which are
section headers, which are spacers, which is the signature row.

Verified by hand against the real bundled templates (see the plan's
"Reference: verified template structure" section) — these counts and keys
are not guesses.
"""

import docx
import pytest

from carescribe.core import clinical_forms


def _load(name):
    return docx.Document(clinical_forms.TEMPLATES_DIR / name)


def test_slugify_collapses_punctuation_and_case():
    assert clinical_forms.slugify("Current functioning") == "current_functioning"
    assert clinical_forms.slugify("Item code (if relevant):") == "item_code_if_relevant"
    assert clinical_forms.slugify("  Mood  ") == "mood"


def test_session_notes_field_walk_finds_nine_fields():
    doc = _load("client_session_notes.docx")
    fields = clinical_forms._walk_table(doc.tables[0], table_index=0, start_row=6)
    assert len(fields) == 9
    keys = [f.key for f in fields]
    assert "session_summary.mental_state_symptoms_if_applicable" in keys
    assert "session_summary.homework_set_and_reviewed" in keys
    # "Review dates:" already has 4 blank trailing paragraphs baked into
    # its own cell (unlike Treatment Review's, which is a bare label whose
    # answer lives in the next blank row) — an own-cell field, not the
    # lookahead case.
    review = next(f for f in fields if f.key == "session_summary.review_dates")
    assert review.value_row_index == 16
    assert review.append_after_label is True


def test_session_notes_signature_row_is_excluded():
    doc = _load("client_session_notes.docx")
    fields = clinical_forms._walk_table(doc.tables[0], table_index=0, start_row=6)
    assert all("signature" not in f.key for f in fields)


def test_treatment_review_spec_has_fourteen_fields():
    spec = clinical_forms.get_form_spec("client_treatment_review")
    assert spec.title == "Client Treatment Review"
    assert len(spec.fields) == 14
    keys = [f.key for f in spec.fields]
    assert "clinical_formulation.current_diagnoses_prognoses" in keys
    # "Review dates:" is a bare label whose own paragraph count is 1 and
    # whose answer lives in the following blank row (row 25).
    review = next(f for f in spec.fields if f.key.endswith("review_dates"))
    assert review.value_row_index == 25


def test_treatment_review_header_fields():
    spec = clinical_forms.get_form_spec("client_treatment_review")
    header_keys = [h.key for h in spec.header_fields]
    assert header_keys == [
        "date", "practitioner", "client_name", "client_dob",
        "session_number", "item_code", "reason_for_referral",
    ]
    reason = next(h for h in spec.header_fields if h.key == "reason_for_referral")
    assert reason.style == "append"
    date = next(h for h in spec.header_fields if h.key == "date")
    assert date.style == "inline"


def test_unknown_form_id_raises():
    with pytest.raises(clinical_forms.ClinicalFormError):
        clinical_forms.get_form_spec("not_a_real_form")
