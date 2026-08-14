"""The three bundled APS templates load and match the structure this
feature's extraction code assumes. If a future template edit changes row
counts, this fails loudly here rather than silently filling the wrong cell.
"""

from pathlib import Path

import docx
import pytest

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "carescribe" / "templates"


@pytest.mark.parametrize(
    "filename,expected_table_count,expected_row_counts",
    [
        ("client_treatment_review.docx", 1, [27]),
        ("client_session_notes.docx", 1, [19]),
        ("biopsychosocial_assessment.docx", 2, [62, 14]),
    ],
)
def test_bundled_template_shape(filename, expected_table_count, expected_row_counts):
    path = TEMPLATES_DIR / filename
    assert path.is_file(), f"missing bundled template: {path}"
    doc = docx.Document(path)
    assert len(doc.tables) == expected_table_count
    assert [len(t.rows) for t in doc.tables] == expected_row_counts
