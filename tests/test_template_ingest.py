"""A clinic can add its own table-based .docx form. The generic parser must
infer everything the three bundled specs hand-code (body start row, header
cell coordinates, the formulation grid range) — proven by reproducing those
specs field-for-field — then persist an uploaded template so it shows up in
the form selector.
"""

import io

import docx
import pytest

from carescribe.core import clinical_forms, template_ingest

BUNDLED = ["client_treatment_review", "client_session_notes", "biopsychosocial_assessment"]


def _anchors(fields):
    return [
        (f.key, f.table_index, f.value_row_index, f.value_col_index, f.append_after_label)
        for f in fields
    ]


@pytest.mark.parametrize("form_id", BUNDLED)
def test_generic_parser_reproduces_bundled_spec(form_id):
    hand = clinical_forms.get_form_spec(form_id)
    generic = template_ingest.parse_template_path(
        clinical_forms.TEMPLATES_DIR / f"{form_id}.docx", form_id=form_id, title=hand.title
    )
    assert _anchors(generic.fields) == _anchors(hand.fields)
    assert [(h.key, h.style, h.table_index, h.row_index, h.col_index) for h in generic.header_fields] == [
        (h.key, h.style, h.table_index, h.row_index, h.col_index) for h in hand.header_fields
    ]


@pytest.mark.parametrize(
    "form_id,expected_body_start",
    [("client_treatment_review", 6), ("client_session_notes", 6), ("biopsychosocial_assessment", 5)],
)
def test_infer_header_finds_body_start(form_id, expected_body_start):
    doc = docx.Document(clinical_forms.TEMPLATES_DIR / f"{form_id}.docx")
    _headers, body_start = template_ingest._infer_header(doc.tables[0])
    assert body_start == expected_body_start


def test_find_grids_detects_the_formulation_grid():
    doc = docx.Document(clinical_forms.TEMPLATES_DIR / "biopsychosocial_assessment.docx")
    grids = template_ingest._find_grids(doc.tables[1], 0)
    assert len(grids) == 1
    banner_row, header_row, first_data_row, last_data_row, section = grids[0]
    assert (header_row, first_data_row, last_data_row) == (1, 2, 5)
    assert section == "CLINICAL FORMULATION"


def test_walk_table_respects_end_row():
    doc = docx.Document(clinical_forms.TEMPLATES_DIR / "biopsychosocial_assessment.docx")
    full = clinical_forms._walk_table(doc.tables[0], table_index=0, start_row=5)
    clipped = clinical_forms._walk_table(doc.tables[0], table_index=0, start_row=5, end_row=17)
    assert len(clipped) < len(full)
    assert clipped == full[: len(clipped)]


# --- synthetic template ------------------------------------------------------

def _merge_full_width(row):
    row.cells[0].merge(row.cells[-1])


def _build_synthetic() -> bytes:
    d = docx.Document()
    d.add_paragraph("Bright Clinic Session Form")
    t = d.add_table(rows=0, cols=4)

    meta = t.add_row().cells
    meta[0].text = "Date:"
    meta[1].text = "Practitioner:"
    meta2 = t.add_row().cells
    meta2[0].text = "Client name:"
    meta2[1].text = "Client DOB:"

    _merge_full_width(t.add_row())  # spacer
    rr = t.add_row()
    _merge_full_width(rr)
    rr.cells[0].text = "Reason for referral:"
    _merge_full_width(t.add_row())  # spacer

    banner = t.add_row()
    _merge_full_width(banner)
    banner.cells[0].text = "SESSION SUMMARY"

    for label in ("Presenting issue", "Intervention applied", "Plan"):
        r = t.add_row()
        r.cells[0].text = label
        r.cells[1].merge(r.cells[-1])

    sig = t.add_row()
    sig.cells[0].text = "Signature:"
    sig.cells[1].merge(sig.cells[-1])

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_parse_minimal_synthetic_template():
    spec = template_ingest.parse_template_bytes(
        _build_synthetic(), form_id="bright_clinic", title="Bright Clinic Session Form"
    )
    assert [h.key for h in spec.header_fields] == [
        "date", "practitioner", "client_name", "client_dob", "reason_for_referral",
    ]
    assert next(h for h in spec.header_fields if h.key == "reason_for_referral").style == "append"
    assert [f.label for f in spec.fields] == [
        "Presenting issue", "Intervention applied", "Plan",
    ]
    assert all(f.key.startswith("session_summary.") for f in spec.fields)
    assert all("signature" not in f.key for f in spec.fields)


def test_round_trip_fill_of_a_parsed_user_template():
    data = _build_synthetic()
    spec = template_ingest.parse_template_bytes(data, form_id="bright_clinic", title="Bright Clinic")
    values = {f.key: f"SENTINEL_{i}" for i, f in enumerate(spec.fields)}
    out = template_ingest.fill_parsed_template(spec, data, values, {"date": "2026-08-29"})
    reopened = docx.Document(io.BytesIO(out))
    table = reopened.tables[0]
    for i, f in enumerate(spec.fields):
        cell = clinical_forms._dedupe_row(table.rows[f.value_row_index])[f.value_col_index]
        assert f"SENTINEL_{i}" in cell.text
    banner_texts = [clinical_forms._dedupe_row(r)[0].text for r in table.rows]
    assert "SESSION SUMMARY" in banner_texts


# --- persistence + registry ------------------------------------------------

@pytest.fixture()
def user_templates(tmp_path, monkeypatch):
    monkeypatch.setattr(template_ingest.desktop, "app_data_dir", lambda: tmp_path)
    clinical_forms.get_form_spec.cache_clear()
    yield tmp_path
    clinical_forms.get_form_spec.cache_clear()


def test_save_lists_and_loads_a_user_template(user_templates):
    form_id = template_ingest.save_template(_build_synthetic(), "Bright Clinic Session Form.docx")
    assert form_id == "bright_clinic_session_form"

    options = dict(clinical_forms.available_forms())
    assert options.get(form_id) == "Bright Clinic Session Form"

    spec = clinical_forms.get_form_spec(form_id)
    assert spec.title == "Bright Clinic Session Form"
    assert [f.label for f in spec.fields] == ["Presenting issue", "Intervention applied", "Plan"]

    # fill_template must find the saved asset on disk
    out = clinical_forms.fill_template(spec, {spec.fields[0].key: "ok"}, {})
    assert docx.Document(io.BytesIO(out))


def test_save_rejects_a_docx_with_no_table(user_templates):
    d = docx.Document()
    d.add_paragraph("Just a letter, no form here.")
    buf = io.BytesIO()
    d.save(buf)
    with pytest.raises(clinical_forms.ClinicalFormError):
        template_ingest.save_template(buf.getvalue(), "letter.docx")


def test_duplicate_upload_name_gets_a_distinct_id(user_templates):
    a = template_ingest.save_template(_build_synthetic(), "Form.docx")
    b = template_ingest.save_template(_build_synthetic(), "Form.docx")
    assert a != b
    assert {a, b} == {"form", "form_2"}


def test_available_forms_survives_a_missing_templates_dir(tmp_path, monkeypatch):
    monkeypatch.setattr(template_ingest.desktop, "app_data_dir", lambda: tmp_path / "nope")
    clinical_forms.get_form_spec.cache_clear()
    ids = [fid for fid, _ in clinical_forms.available_forms()]
    assert ids == BUNDLED


def test_a_user_template_is_a_first_class_citizen_in_generation(user_templates):
    form_id = template_ingest.save_template(_build_synthetic(), "Bright Clinic.docx")
    spec = clinical_forms.get_form_spec(form_id)

    system, user = clinical_forms.build_prompt(spec, "[PATIENT] attended. Plan: review in a week.")
    for field in spec.fields:
        assert f"<<FIELD:{field.key}>>" in system
    assert "[PATIENT]" in user

    raw = "".join(f"<<FIELD:{f.key}>>\ntext {i}\n" for i, f in enumerate(spec.fields))
    parsed = clinical_forms.parse_fields(spec, raw)
    assert parsed == {f.key: f"text {i}" for i, f in enumerate(spec.fields)}
