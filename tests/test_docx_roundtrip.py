"""
Word round-trip: upload -> redact -> download, structure preserved.

The document under test is built in code rather than checked in, so the awkward
shapes are explicit: a details table, a name Word fragmented across runs, a name
wrapped onto the next paragraph, and identifiers in the header and footer.

Everything here is fabricated.
"""

import io

import pytest

from docx import Document as DocxDocument

from carescribe.core import batch, deidentify, docx_redact

PATIENT = "Oluwaseun Adeyinka"
NHS = "943 476 5919"
WARD = "Cedar Ward"

APPROVED_MAP = {
    PATIENT: "[PATIENT]",
    "Oluwaseun": "[PATIENT]",
    "Adeyinka": "[PATIENT]",
    NHS: "[NHS_NO]",
    WARD: "[WARD]",
    "990214": "[MRN]",
}


def _build(path, *, with_text_box: bool = False):
    doc = DocxDocument()
    doc.add_heading("Community Mental Health Team", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    for row, (label, value) in enumerate(
        [("Patient", PATIENT), ("NHS No", NHS), ("Case No", "990214"), ("Ward", WARD)]
    ):
        table.cell(row, 0).text = label
        table.cell(row, 1).text = value

    # A name Word split across runs, as it arrives after an edit.
    fragmented = doc.add_paragraph()
    for piece in ["Olu", "wase", "un Ade", "yinka", " attended the day unit."]:
        fragmented.add_run(piece)

    # A name wrapped onto the following paragraph.
    doc.add_paragraph("The assessment was completed by Oluwaseun")
    doc.add_paragraph("Adeyinka in the afternoon.")

    doc.add_paragraph(
        "Detained under Section 2 of the Mental Health Act. HoNOS was 18 and the "
        "PHQ-9 was 12. Olanzapine 10mg at night was continued."
    )

    section = doc.sections[0]
    section.header.paragraphs[0].text = f"{PATIENT} — NHS {NHS}"
    section.footer.paragraphs[0].text = f"{WARD} — confidential"

    if with_text_box:
        # A minimal text box. Its content lives outside the paragraph tree, which
        # is exactly why the redaction pass cannot reach it.
        doc.add_paragraph().add_run()._r.append(_text_box_xml())

    doc.save(str(path))
    return path


_TEXT_BOX_XML = """
<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml">
  <v:shape><v:textbox><w:txbxContent>
    <w:p><w:r><w:t>{text}</w:t></w:r></w:p>
  </w:txbxContent></v:textbox></v:shape>
</w:pict>
"""


def _text_box_xml():
    from docx.oxml import parse_xml

    return parse_xml(_TEXT_BOX_XML.format(text=f"{PATIENT} is only in this box."))


@pytest.fixture(scope="module")
def redacted(tmp_path_factory):
    folder = tmp_path_factory.mktemp("docx")
    source = _build(folder / "referral.docx")
    output = folder / "referral.deid.docx"
    docx_redact.apply_redactions(str(source), str(output), APPROVED_MAP)
    return output


# ==========================================================================
# (a) every identifier is gone
# ==========================================================================

@pytest.mark.parametrize(
    "value", [PATIENT, "Oluwaseun", "Adeyinka", NHS, WARD, "990214"]
)
def test_identifiers_are_gone(redacted, value):
    assert value not in docx_redact.extract_text(str(redacted))


def test_the_run_fragmented_name_is_redacted(redacted):
    text = docx_redact.extract_text(str(redacted))
    assert "[PATIENT] attended the day unit." in text


def test_the_paragraph_wrapped_name_is_redacted(redacted):
    text = docx_redact.extract_text(str(redacted))
    assert "completed by [PATIENT]" in text
    assert "in the afternoon." in text


# ==========================================================================
# (b) the table survives intact
# ==========================================================================

def test_the_table_keeps_its_shape_and_style(redacted):
    doc = DocxDocument(str(redacted))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 4
    assert len(table.columns) == 2
    assert table.style.name == "Table Grid"


def test_the_table_labels_are_untouched(redacted):
    table = DocxDocument(str(redacted)).tables[0]
    assert [table.cell(r, 0).text for r in range(4)] == [
        "Patient", "NHS No", "Case No", "Ward"
    ]


def test_the_table_values_are_placeholders(redacted):
    table = DocxDocument(str(redacted)).tables[0]
    assert [table.cell(r, 1).text for r in range(4)] == [
        "[PATIENT]", "[NHS_NO]", "[MRN]", "[WARD]"
    ]


# ==========================================================================
# (c) header and footer keep everything but the redacted spans
# ==========================================================================

def test_header_and_footer_are_redacted_but_otherwise_preserved(redacted):
    section = DocxDocument(str(redacted)).sections[0]
    header = section.header.paragraphs[0].text
    footer = section.footer.paragraphs[0].text
    assert header == "[PATIENT] — NHS [NHS_NO]"
    assert footer == "[WARD] — confidential"


def test_the_heading_survives(redacted):
    assert "Community Mental Health Team" in docx_redact.extract_text(str(redacted))


# ==========================================================================
# (d) precision terms survive
# ==========================================================================

@pytest.mark.parametrize(
    "term",
    ["Section 2", "Mental Health Act", "HoNOS", "PHQ-9", "Olanzapine 10mg",
     "at night", "the day unit"],
)
def test_clinical_content_survives(redacted, term):
    assert term in docx_redact.extract_text(str(redacted))


# ==========================================================================
# (e) the safety sweep on the Word output is clean
# ==========================================================================

def test_residual_scan_on_the_word_output_is_empty(redacted):
    assert deidentify.residual_scan(docx_redact.extract_text(str(redacted))) == []


# ==========================================================================
# The write path and its gate
# ==========================================================================

def test_write_approved_docx_writes_and_sweeps(tmp_path, monkeypatch):
    monkeypatch.setattr(batch, "OUTPUT_DIR", tmp_path / "out")
    source = _build(tmp_path / "referral.docx")
    written = batch.write_approved_docx(
        "referral.docx", source.read_bytes(), APPROVED_MAP
    )
    assert written.exists()
    assert written.name.endswith(batch.APPROVED_DOCX_SUFFIX)
    assert deidentify.residual_scan(docx_redact.extract_text(str(written))) == []


def test_write_approved_docx_refuses_when_the_map_misses_an_identifier(
    tmp_path, monkeypatch
):
    """The Word path must clear the same bar as the text path."""
    monkeypatch.setattr(batch, "OUTPUT_DIR", tmp_path / "out")
    source = _build(tmp_path / "referral.docx")
    leaky = {k: v for k, v in APPROVED_MAP.items() if k != NHS}
    with pytest.raises(batch.BatchError) as excinfo:
        batch.write_approved_docx("referral.docx", source.read_bytes(), leaky)
    assert "Refusing to write" in str(excinfo.value)
    assert not batch.approved_docx_path("referral.docx").exists()


def test_write_approved_docx_refuses_when_a_raw_cr_hides_an_identifier(
    tmp_path, monkeypatch
):
    """The residual safety-net scan must not be foolable by a stray \\r.

    A run's text can carry a literal CR (0x0D) in its XML — a real character,
    not a paragraph break, and something non-Word tooling can produce. Word
    itself normalises this on open, but python-docx (and this app's own
    reader) will happily read it back verbatim. Before ingest.py's
    normalisation was threaded through this call site too, a letterhead-style
    "Town, County" line sitting right after such a CR was invisible to the
    residual scan's \\n-anchored HEADER_LOCATION pattern — the exact
    under-redaction this test exists to catch before it ever reaches disk.
    """
    from docx.oxml import parse_xml

    monkeypatch.setattr(batch, "OUTPUT_DIR", tmp_path / "out")
    source_path = _build(tmp_path / "referral.docx")

    doc = DocxDocument(str(source_path))
    paragraph = doc.add_paragraph()
    run_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:t xml:space="preserve">Wetherby, West Yorkshire&#13;'
        "an unrelated closing line</w:t></w:r>"
    )
    paragraph._p.append(parse_xml(run_xml))
    doc.save(str(source_path))

    # "Wetherby, West Yorkshire" is deliberately absent from the approved map
    # — a value the reviewer's map did not cover, same shape as the sibling
    # test above — so it is left untouched by redaction and must be caught by
    # the residual scan, CR or not.
    with pytest.raises(batch.BatchError) as excinfo:
        batch.write_approved_docx(
            "referral.docx", source_path.read_bytes(), APPROVED_MAP
        )
    assert "Refusing to write" in str(excinfo.value)
    assert not batch.approved_docx_path("referral.docx").exists()


def test_the_original_document_is_never_written_to_disk(tmp_path, monkeypatch):
    """Redaction happens in memory — the un-redacted file must not be staged.

    Staging the source through a temp file would put the PHI-bearing original on
    disk, briefly, and not at all if the process died mid-way. The invariant is
    that the only bytes this app ever writes are de-identified.
    """
    monkeypatch.setattr(batch, "OUTPUT_DIR", tmp_path / "out")
    source = _build(tmp_path / "referral.docx")
    original = source.read_bytes()

    written: list[tuple[str, bytes]] = []
    real_write_bytes = type(source).write_bytes

    def spy(self, data):
        written.append((str(self), bytes(data)))
        return real_write_bytes(self, data)

    monkeypatch.setattr(type(source), "write_bytes", spy)
    batch.write_approved_docx("referral.docx", original, APPROVED_MAP)

    assert written, "nothing was written at all"
    for path, data in written:
        assert data != original, f"the original document was written to {path}"
        assert PATIENT not in docx_redact.extract_text(io.BytesIO(data))


def test_write_approved_docx_needs_the_original(tmp_path, monkeypatch):
    monkeypatch.setattr(batch, "OUTPUT_DIR", tmp_path / "out")
    with pytest.raises(batch.BatchError):
        batch.write_approved_docx("referral.docx", b"", APPROVED_MAP)


# ==========================================================================
# B3 — the approved map is built from the approved table
# ==========================================================================

def test_the_approved_map_covers_every_surface_form():
    entities = [
        {"type": "PATIENT_NAME", "value": "Mariam Aisha Rahman",
         "placeholder": "[PATIENT]", "action": "Redact"},
    ]
    literals = batch.approved_map(entities, known_as="Mimi")
    assert literals["Mariam Aisha Rahman"] == "[PATIENT]"
    for form in ("Mariam", "Rahman", "Mrs Rahman", "M.A.R.", "Mimi"):
        assert literals.get(form) == "[PATIENT]", form


def test_a_kept_row_contributes_nothing_to_the_approved_map():
    entities = [
        {"type": "PATIENT_NAME", "value": "Wei Chen",
         "placeholder": "[PATIENT]", "action": "Keep"},
        {"type": "MRN", "value": "990214", "placeholder": "[MRN]", "action": "Redact"},
    ]
    literals = batch.approved_map(entities)
    assert "Wei Chen" not in literals
    assert literals["990214"] == "[MRN]"


def test_the_approved_map_is_longest_literal_first():
    entities = [
        {"type": "PATIENT_NAME", "value": "Wei Chen",
         "placeholder": "[PATIENT]", "action": "Redact"},
    ]
    lengths = [len(k) for k in batch.approved_map(entities)]
    assert lengths == sorted(lengths, reverse=True)


# ==========================================================================
# B6 — text boxes are surfaced, never silently passed
# ==========================================================================

def test_a_document_with_a_text_box_is_flagged(tmp_path):
    source = _build(tmp_path / "boxed.docx", with_text_box=True)
    assert batch.document_has_text_boxes(source.read_bytes())


def test_a_plain_document_is_not_flagged(tmp_path):
    source = _build(tmp_path / "plain.docx")
    assert not batch.document_has_text_boxes(source.read_bytes())


def test_loading_a_docx_keeps_the_original_and_flags_text_boxes(tmp_path):
    plain = _build(tmp_path / "plain.docx")
    boxed = _build(tmp_path / "boxed.docx", with_text_box=True)
    documents, errors = batch.load_documents([str(plain), str(boxed)])
    assert not errors
    assert documents["plain.docx"].source_bytes
    assert not documents["plain.docx"].has_text_boxes
    assert documents["boxed.docx"].has_text_boxes


def test_the_full_upload_redact_download_round_trip(tmp_path, monkeypatch):
    """Load -> detect -> approve -> Word out, using only the approved map."""
    monkeypatch.setattr(batch, "OUTPUT_DIR", tmp_path / "out")
    source = _build(tmp_path / "referral.docx")

    documents, errors = batch.load_documents([str(source)])
    assert not errors
    document = batch.analyze_document(documents["referral.docx"])

    written = batch.write_approved_docx(
        document.name,
        document.source_bytes,
        batch.approved_map(document.entities, document.known_as),
    )

    text = docx_redact.extract_text(str(written))
    assert PATIENT not in text
    assert NHS not in text
    assert deidentify.residual_scan(text) == []
    # Structure and clinical meaning both intact.
    assert len(DocxDocument(str(written)).tables[0].rows) == 4
    assert "Mental Health Act" in text


def test_text_box_content_really_does_survive_redaction(tmp_path):
    """The warning exists because this is true — assert it, don't assume it."""
    source = _build(tmp_path / "boxed.docx", with_text_box=True)
    output = tmp_path / "boxed.deid.docx"
    docx_redact.apply_redactions(str(source), str(output), APPROVED_MAP)
    from docx import Document as D
    assert PATIENT in D(str(output)).element.body.xml
