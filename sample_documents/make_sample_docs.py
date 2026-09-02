"""Generate synthetic, complex .docx source documents for manually testing
the clinical-forms pipeline (upload -> de-identify -> combine -> generate).

All names/dates/numbers are fabricated. Nothing here is real PHI.
"""
from __future__ import annotations

import os

from docx import Document

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(OUT_DIR, exist_ok=True)


def _heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _para(doc, text, bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        run = p.add_run(bold_label)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def _two_col_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    return table


def _grid_table(doc, header, rows):
    table = doc.add_table(rows=0, cols=len(header))
    table.style = "Table Grid"
    hdr = table.add_row()
    for i, h in enumerate(header):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_vals in rows:
        row = table.add_row()
        for i, v in enumerate(row_vals):
            row.cells[i].text = v
    return table


# ---------------------------------------------------------------------------
# Document 1: GP referral letter
# ---------------------------------------------------------------------------
def build_referral_letter(path):
    doc = Document()
    doc.add_heading("Referral Letter — Riverside Family Medical Practice", level=0)

    _para(doc, "12 Wattle Street, Coburg VIC 3058  |  Ph: (03) 5551 0192  |  Fax: (03) 5551 0193")
    doc.add_paragraph()

    _para(doc, "3 February 2026", bold_label="Date: ")
    _para(doc, "Dr. Amelia Ferro, Clinical Psychologist, Northgate Psychology Clinic",
          bold_label="To: ")
    doc.add_paragraph()

    _heading(doc, "Patient Details", level=2)
    _two_col_table(doc, [
        ("Full name", "Jordan Elliot Whitfield"),
        ("Date of birth", "12/04/1985"),
        ("Medicare number", "2934 5671 0"),
        ("UR / Hospital No.", "RFMP-88213"),
        ("Address", "45 Kestrel Ave, Coburg VIC 3058"),
        ("Phone", "0412 887 234"),
        ("Email", "jordan.whitfield85@example.com"),
        ("Next of kin", "Priya Whitfield (spouse) — 0433 990 214"),
        ("Referring GP", "Dr. Susan Ng, Riverside Family Medical Practice"),
        ("Referral date", "3 February 2026"),
        ("Medicare rebate item", "Mental Health Care Plan — Item 2710, sessions 1–6"),
    ])
    doc.add_paragraph()

    _heading(doc, "Reason for Referral", level=2)
    doc.add_paragraph(
        "Mr Whitfield is a 40-year-old warehouse supervisor referred for psychological "
        "assessment and treatment of persistent low mood, anxiety and sleep disturbance "
        "following a workplace incident in October 2025 (forklift near-miss involving a "
        "colleague). He reports intrusive memories of the event, hypervigilance around "
        "machinery, and has been avoiding the warehouse floor where possible. Symptoms "
        "have been present for approximately 4 months and are affecting his ability to "
        "perform his usual duties and his relationship with his partner."
    )

    _heading(doc, "Current Medications", level=2)
    _grid_table(doc, ["Medication", "Dose", "Frequency", "Prescriber", "Start date"], [
        ["Sertraline", "50 mg", "Once daily (morning)", "Dr. Susan Ng", "18/11/2025"],
        ["Temazepam", "10 mg", "PRN, night, max 2x/week", "Dr. Susan Ng", "18/11/2025"],
        ["Atorvastatin", "20 mg", "Once daily (evening)", "Dr. Susan Ng", "2021 (ongoing)"],
    ])
    doc.add_paragraph()

    _heading(doc, "Past Medical & Psychiatric History", level=2)
    doc.add_paragraph(
        "No prior psychiatric admissions or treatment. Hypercholesterolaemia, diagnosed "
        "2021, managed with statin. Appendicectomy 2009. No known drug allergies. "
        "Non-smoker; reports 6–8 standard drinks per week, increased from ~2/week prior "
        "to the workplace incident."
    )

    _heading(doc, "Family History", level=2)
    _grid_table(doc, ["Relationship", "Relevant history"], [
        ["Mother", "Generalised anxiety disorder, managed in community"],
        ["Father", "Nil known psychiatric history; hypertension"],
        ["Sibling (younger brother)", "Nil known"],
    ])

    doc.add_paragraph()
    _para(doc, "Please assess and manage as appropriate. Happy to discuss — my direct line "
                "is (03) 5551 0192 ext. 4.", bold_label="")
    doc.add_paragraph()
    _para(doc, "Dr. Susan Ng, MBBS FRACGP", bold_label="")
    _para(doc, "Provider No. 2481726A", bold_label="")

    doc.save(path)


# ---------------------------------------------------------------------------
# Document 2: Biopsychosocial intake assessment notes
# ---------------------------------------------------------------------------
def build_intake_notes(path):
    doc = Document()
    doc.add_heading("Initial Assessment / Intake Notes", level=0)
    _two_col_table(doc, [
        ("Client", "Jordan Elliot Whitfield"),
        ("DOB", "12/04/1985"),
        ("Assessment date", "10 February 2026"),
        ("Clinician", "Dr. Amelia Ferro, Clinical Psychologist"),
        ("Session type", "Initial consultation (90 min)"),
    ])
    doc.add_paragraph()

    _heading(doc, "Presenting Problem", level=2)
    doc.add_paragraph(
        "Client presents with symptoms consistent with an adjustment disorder with "
        "mixed anxiety and depressed mood, precipitated by a workplace near-miss "
        "incident in October 2025. Reports intrusive recollections 3–4 times per week, "
        "sleep onset latency of 60+ minutes, early morning waking, decreased appetite, "
        "and irritability noted by his partner. Denies suicidal ideation, denies "
        "self-harm history. Some passive thoughts of 'not wanting to wake up' during the "
        "worst week (early December 2025), no plan or intent, resolved with improved sleep."
    )

    _heading(doc, "Developmental & Personal History", level=2)
    doc.add_paragraph(
        "Born in Bendigo VIC, raised by both biological parents alongside one younger "
        "sibling. Reports a stable, unremarkable childhood with no history of abuse or "
        "neglect disclosed. Completed Year 12, followed by a Certificate III in Logistics. "
        "Describes himself premorbidly as even-tempered, socially engaged, and physically "
        "active (plays weekend football)."
    )

    _heading(doc, "Family History", level=2)
    _grid_table(doc, ["Relationship", "Name", "Age", "Health / notes"], [
        ["Spouse", "Priya Whitfield", "38", "Registered nurse; supportive, some strain reported"],
        ["Son", "Kai Whitfield", "9", "No concerns"],
        ["Daughter", "Mira Whitfield", "6", "No concerns"],
        ["Mother", "Denise Whitfield", "68", "GAD, managed in community by GP"],
        ["Father", "Robert Whitfield", "71", "Hypertension; otherwise well"],
    ])
    doc.add_paragraph()

    _heading(doc, "Social & Occupational History", level=2)
    doc.add_paragraph(
        "Employed as a warehouse supervisor at Coburg Logistics Pty Ltd for 11 years, "
        "currently on modified duties (office-based) since the incident. Married 12 years. "
        "Home owner. Reports a small but reliable social network — weekend football club "
        "teammates and two close friends from high school. Financial stress described as "
        "'manageable but tight' due to reduced overtime since moving to modified duties."
    )

    _heading(doc, "Substance Use", level=2)
    _grid_table(doc, ["Substance", "Pattern", "Change since incident"], [
        ["Alcohol", "6–8 standard drinks/week, mostly weekends", "Increased from ~2/week"],
        ["Tobacco", "Non-smoker", "N/A"],
        ["Caffeine", "3–4 coffees/day", "Increased from 1–2/day"],
        ["Illicit drugs", "Denies current or past use", "N/A"],
    ])
    doc.add_paragraph()

    _heading(doc, "Risk Assessment", level=2)
    _grid_table(doc, ["Domain", "Level", "Notes"], [
        ["Suicidal ideation", "Low", "Passive ideation Dec 2025, resolved; no current SI"],
        ["Self-harm", "Low", "No history"],
        ["Risk to others", "Nil identified", "No history of aggression"],
        ["Risk from others", "Nil identified", "Denies family violence, denies coercion"],
        ["Occupational risk", "Moderate", "Hypervigilance around machinery on warehouse floor"],
    ])
    doc.add_paragraph()

    _heading(doc, "Mental State on Presentation", level=2)
    doc.add_paragraph(
        "Appearance and behaviour: casually dressed, good eye contact, mild psychomotor "
        "tension. Mood: 'flat, on edge.' Affect: reactive, congruent, mildly restricted "
        "range. Speech: normal rate and volume. Thought form: linear, goal-directed. "
        "Thought content: no delusions; preoccupation with safety at work. Perception: no "
        "hallucinations reported. Cognition: alert, oriented x4, attention intact. Insight: "
        "good. Judgement: intact."
    )

    _heading(doc, "Formulation Notes (clinician working notes)", level=2)
    doc.add_paragraph(
        "Predisposing: family history of anxiety, high personal standards at work. "
        "Precipitating: October 2025 forklift near-miss involving a colleague. "
        "Perpetuating: avoidance of warehouse floor, increased alcohol/caffeine use, "
        "sleep disruption maintaining hyperarousal. Protective: stable marriage, supportive "
        "spouse, engaged employer offering modified duties, no prior psychiatric history, "
        "good insight and motivation for treatment."
    )

    doc.save(path)


# ---------------------------------------------------------------------------
# Document 3: Session log + treatment progress notes
# ---------------------------------------------------------------------------
def build_session_log(path):
    doc = Document()
    doc.add_heading("Session Log & Progress Notes", level=0)
    _two_col_table(doc, [
        ("Client", "Jordan Elliot Whitfield"),
        ("Treating clinician", "Dr. Amelia Ferro, Clinical Psychologist"),
        ("Treatment modality", "Trauma-focused CBT with graded exposure"),
        ("Referral basis", "Mental Health Care Plan, Item 2710"),
    ])
    doc.add_paragraph()

    _heading(doc, "Session Summary Table", level=2)
    _grid_table(
        doc,
        ["Date", "Session #", "Duration", "Format", "Attendance", "Summary"],
        [
            ["10/02/2026", "1", "90 min", "In person", "Attended", "Intake & biopsychosocial assessment completed."],
            ["17/02/2026", "2", "50 min", "In person", "Attended", "Psychoeducation on trauma response; sleep hygiene plan set."],
            ["24/02/2026", "3", "50 min", "Telehealth", "Attended", "Introduced diaphragmatic breathing; reviewed sleep diary — modest improvement."],
            ["03/03/2026", "4", "50 min", "In person", "Attended", "Began graded exposure hierarchy for warehouse floor exposure."],
            ["10/03/2026", "5", "50 min", "In person", "Cancelled — client unwell", "N/A"],
            ["17/03/2026", "6", "50 min", "In person", "Attended", "Exposure step 1 completed (observed floor from office window) with reducing SUDS 8→4."],
        ],
    )
    doc.add_paragraph()

    _heading(doc, "Standardised Measures", level=2)
    _grid_table(doc, ["Measure", "Date", "Score", "Interpretation"], [
        ["DASS-21 (Depression)", "10/02/2026", "18", "Moderate"],
        ["DASS-21 (Anxiety)", "10/02/2026", "16", "Moderate"],
        ["DASS-21 (Stress)", "10/02/2026", "22", "Severe"],
        ["DASS-21 (Depression)", "17/03/2026", "10", "Mild"],
        ["DASS-21 (Anxiety)", "17/03/2026", "9", "Mild"],
        ["DASS-21 (Stress)", "17/03/2026", "12", "Moderate"],
        ["PCL-5 (PTSD checklist)", "10/02/2026", "39", "Above clinical cutoff"],
        ["PCL-5 (PTSD checklist)", "17/03/2026", "24", "Below clinical cutoff"],
    ])
    doc.add_paragraph()

    _heading(doc, "Treatment Goals", level=2)
    _grid_table(doc, ["Goal", "Intervention", "Progress"], [
        ["Reduce intrusive memories to <1/week", "Trauma-focused CBT, imaginal exposure", "Reduced from 3–4/week to ~1/week"],
        ["Return to full warehouse-floor duties", "Graded in-vivo exposure hierarchy", "Step 1 of 6 completed"],
        ["Improve sleep onset to <30 min", "Sleep hygiene, stimulus control, reduced evening caffeine", "Sleep onset improved to ~35 min"],
        ["Reduce alcohol to <4 standard drinks/week", "Motivational interviewing, self-monitoring diary", "Reduced to ~5/week"],
    ])
    doc.add_paragraph()

    _heading(doc, "Clinician Impression", level=2)
    doc.add_paragraph(
        "Client is engaging well with treatment and demonstrating a clear reduction in "
        "trauma-related symptoms across both standardised measures and session-report. "
        "Continues to benefit from graded exposure; plan to progress to step 2 of the "
        "warehouse-floor hierarchy (brief supervised entry with a colleague present) at "
        "next session. No risk concerns identified at this time. Recommend continuing "
        "under the current Mental Health Care Plan with a review letter to the referring "
        "GP prior to session 6 rebate exhaustion."
    )

    doc.save(path)


# ---------------------------------------------------------------------------
# Document 4: Treatment plan review letter (for the Treatment Review form)
# ---------------------------------------------------------------------------
def build_treatment_review_source(path):
    doc = Document()
    doc.add_heading("Treatment Plan Review — Correspondence to Referrer", level=0)
    _two_col_table(doc, [
        ("Client", "Jordan Elliot Whitfield"),
        ("DOB", "12/04/1985"),
        ("Review date", "17 March 2026"),
        ("Sessions completed", "5 of 6 (1 cancelled, rebooked)"),
        ("Referring GP", "Dr. Susan Ng"),
        ("Next review due", "After session 6, prior to further MHCP referral"),
    ])
    doc.add_paragraph()

    _heading(doc, "Summary of Treatment to Date", level=2)
    doc.add_paragraph(
        "Mr Whitfield has attended 5 of 6 funded sessions under his current Mental Health "
        "Care Plan, engaging in trauma-focused CBT with graded exposure for an adjustment "
        "disorder with mixed anxiety and depressed mood following a workplace near-miss "
        "in October 2025. Treatment has focused on psychoeducation, sleep hygiene, "
        "diaphragmatic breathing, and a graded in-vivo exposure hierarchy targeting "
        "avoidance of the warehouse floor."
    )

    _heading(doc, "Outcome Measures Comparison", level=2)
    _grid_table(doc, ["Measure", "Baseline (10/02/2026)", "Current (17/03/2026)", "Change"], [
        ["DASS-21 Depression", "18 (Moderate)", "10 (Mild)", "Improved"],
        ["DASS-21 Anxiety", "16 (Moderate)", "9 (Mild)", "Improved"],
        ["DASS-21 Stress", "22 (Severe)", "12 (Moderate)", "Improved"],
        ["PCL-5", "39 (above cutoff)", "24 (below cutoff)", "Improved — below clinical cutoff"],
    ])
    doc.add_paragraph()

    _heading(doc, "Risk Review", level=2)
    doc.add_paragraph(
        "No current suicidal ideation, self-harm, or risk to others identified. Passive "
        "ideation reported at intake (December 2025, prior to treatment) has not recurred. "
        "Occupational hypervigilance persists but is reducing in line with exposure progress."
    )

    _heading(doc, "Recommendations", level=2)
    doc.add_paragraph(
        "Recommend a further course of treatment (additional Mental Health Care Plan "
        "sessions) to complete the exposure hierarchy and consolidate gains, with an "
        "anticipated further 4–6 sessions. Client is motivated and prognosis is good given "
        "response to date. Will provide a further review at completion of the exposure "
        "hierarchy or in 8 weeks, whichever is sooner."
    )

    doc.save(path)


# ---------------------------------------------------------------------------
# Document 5: Hospital discharge summary
# ---------------------------------------------------------------------------
def build_discharge_summary(path):
    doc = Document()
    doc.add_heading("Discharge Summary — Merri Creek District Hospital", level=0)

    _para(doc, "210 Sydney Road, Coburg VIC 3058  |  Ph: (03) 5551 7700  |  Fax: (03) 5551 7701")
    doc.add_paragraph()

    _heading(doc, "Patient Details", level=2)
    _two_col_table(doc, [
        ("Full name", "Jordan Elliot Whitfield"),
        ("Date of birth", "12/04/1985"),
        ("Medicare number", "2934 5671 0"),
        ("UR number", "MCDH-410287"),
        ("Address", "45 Kestrel Ave, Coburg VIC 3058"),
        ("Admission date", "14 October 2025"),
        ("Discharge date", "16 October 2025"),
        ("Admitting unit", "Emergency / Orthopaedics (short-stay)"),
        ("Treating consultant", "Dr. Farid Kassab, Orthopaedic Registrar"),
        ("GP copied", "Dr. Susan Ng, Riverside Family Medical Practice"),
    ])
    doc.add_paragraph()

    _heading(doc, "Presenting Complaint", level=2)
    doc.add_paragraph(
        "Mr Whitfield presented to the Emergency Department via ambulance following a workplace "
        "incident at Coburg Logistics Pty Ltd on 14 October 2025, where he stepped back sharply to "
        "avoid a reversing forklift being operated by a colleague and fell against a steel shelving "
        "unit. He sustained a closed injury to the right wrist and superficial lacerations to the "
        "right forearm. No loss of consciousness and no head injury. He was visibly shaken in the "
        "department and reported his heart 'racing' for over an hour afterwards, with intrusive "
        "replaying of the near-miss."
    )

    _heading(doc, "Medications Changed on Discharge", level=2)
    _grid_table(doc, ["Medication", "Prior dose", "New dose", "Reason"], [
        ["Paracetamol", "Nil", "1 g PO QID PRN", "Post-fracture analgesia"],
        ["Oxycodone IR (Endone)", "Nil", "5 mg PO 4–6 hourly PRN, max 5 days", "Short-term breakthrough pain; wean before GP review"],
        ["Naproxen", "Nil", "Not commenced — avoided", "NSAID deferred pending orthopaedic review of fracture healing"],
        ["Atorvastatin", "20 mg nocte", "20 mg nocte (unchanged)", "Continue usual dose — no interaction with new analgesia"],
    ])
    doc.add_paragraph()

    _heading(doc, "Discharge Diagnosis", level=2)
    doc.add_paragraph(
        "1. Closed, minimally displaced fracture of the right distal radius — managed conservatively "
        "with a below-elbow backslab; orthopaedic outpatient review in 7–10 days for cast conversion "
        "and repeat X-ray. 2. Superficial lacerations, right forearm, closed with sutures — for "
        "removal by GP in 7 days. 3. Acute stress reaction following the precipitating workplace "
        "incident — settled with reassurance overnight; no evidence of head injury. GP follow-up "
        "recommended, with consideration of psychological referral if symptoms of anxiety, sleep "
        "disturbance or intrusive memories persist beyond 2–3 weeks."
    )

    _heading(doc, "Follow-up Plan", level=2)
    doc.add_paragraph(
        "Discharged home with partner (Priya Whitfield) on 16 October 2025 in a stable condition, "
        "mobilising independently, right upper limb in backslab with sling. Certified unfit for "
        "usual warehouse duties for a minimum of 4 weeks; suitable for modified/office-based duties "
        "once cleared by GP. Copy of this summary forwarded to Dr. Susan Ng, Riverside Family "
        "Medical Practice, for ongoing care."
    )

    doc.save(path)


# ---------------------------------------------------------------------------
# Document 6: Structured risk assessment
# ---------------------------------------------------------------------------
def build_risk_assessment(path):
    doc = Document()
    doc.add_heading("Structured Risk Assessment — Northgate Psychology Clinic", level=0)

    _para(doc, "8 Derby Street, Pascoe Vale VIC 3044  |  Ph: (03) 5551 3320  |  Fax: (03) 5551 3321")
    doc.add_paragraph()

    _heading(doc, "Client Details", level=2)
    _two_col_table(doc, [
        ("Full name", "Jordan Elliot Whitfield"),
        ("Date of birth", "12/04/1985"),
        ("Medicare number", "2934 5671 0"),
        ("Assessment date", "10 February 2026"),
        ("Assessing clinician", "Dr. Amelia Ferro, Clinical Psychologist"),
        ("Reason for assessment", "Standard risk screen at initial consultation"),
    ])
    doc.add_paragraph()

    _heading(doc, "Risk Domain Ratings", level=2)
    _two_col_table(doc, [
        ("Self-harm", "Low"),
        ("Suicide", "Low"),
        ("Harm to others", "Nil identified"),
        ("Self-neglect", "Low"),
    ])
    doc.add_paragraph()

    _heading(doc, "Self-Harm", level=2)
    doc.add_paragraph(
        "No history of self-harm at any age and no current urges or intent reported. Denies any "
        "current or historical self-injurious behaviour. No identified means or method. Rated low "
        "risk."
    )

    _heading(doc, "Suicide", level=2)
    doc.add_paragraph(
        "Reports brief passive ideation ('not wanting to wake up') during the worst week of "
        "symptoms in early December 2025, with no plan, no intent and no access to means "
        "identified at the time. No history of prior attempts. Ideation resolved alongside "
        "improved sleep and has not recurred; denies any current ideation, plan or intent on "
        "interview today. Protective factors include a stable marriage, supportive spouse, two "
        "young children, an engaged employer, and good insight. Rated low risk."
    )

    _heading(doc, "Harm to Others", level=2)
    doc.add_paragraph(
        "No history of violence, aggression or threatening behaviour. Denies any current thoughts "
        "of harming others and denies access to weapons. Some increased irritability noted by his "
        "partner is assessed as consistent with adjustment symptoms rather than a risk indicator. "
        "Rated nil identified risk."
    )

    _heading(doc, "Self-Neglect", level=2)
    doc.add_paragraph(
        "Self-care, hygiene, nutrition and household responsibilities remain at a normal level. "
        "Continues to attend work on modified duties and manage day-to-day responsibilities. Has "
        "paused his usual weekend football but reports this is by choice rather than inability to "
        "function. No concerns regarding capacity for self-care. Rated low risk."
    )
    doc.add_paragraph()

    _heading(doc, "Safety Plan", level=2)
    _grid_table(doc, ["Element", "Detail"], [
        ["Warning signs", "Sleep onset over 60 minutes for 3+ nights running; increased intrusive memories; withdrawing from Priya and the children"],
        ["Internal coping strategies", "Diaphragmatic breathing (as taught in session); short walk; journalling"],
        ["Social contacts for distraction", "Call brother-in-law; weekend football club group chat; visit parents in Bendigo"],
        ["People who can help", "Priya Whitfield (spouse) — 0433 990 214"],
        ["Professionals / agencies to contact", "Dr. Amelia Ferro (03) 5551 3320; Dr. Susan Ng (03) 5551 0192; Lifeline 13 11 14; Suicide Call Back Service 1300 659 467"],
        ["Means restriction", "No firearms in the home; household medications moved to a locked cabinet at Priya's suggestion as a precaution"],
        ["Reasons for living", "Children Kai and Mira; wanting to return to full duties and football; marriage to Priya"],
    ])
    doc.add_paragraph()

    _heading(doc, "Clinician Summary", level=2)
    doc.add_paragraph(
        "Overall risk is assessed as low across all domains at this time, reflecting a single "
        "resolved episode of passive ideation in December 2025 rather than an ongoing pattern. No "
        "changes to the current outpatient treatment plan are indicated on risk grounds. Safety "
        "plan provided to client and reviewed collaboratively; to be revisited at each session and "
        "immediately if any of the warning signs above recur."
    )

    doc.save(path)


# ---------------------------------------------------------------------------
# Document 7: MDT case conference minutes
# ---------------------------------------------------------------------------
def build_case_conference_note(path):
    doc = Document()
    doc.add_heading("MDT Case Conference — Minutes", level=0)

    _para(doc, "Northgate Psychology Clinic, 8 Derby Street, Pascoe Vale VIC 3044  |  Ph: (03) 5551 3320")
    doc.add_paragraph()

    _para(doc, "24 March 2026, 2:00pm–2:45pm", bold_label="Date/Time: ")
    _para(doc, "Return-to-work planning and continuation of psychological treatment",
          bold_label="Purpose: ")
    _para(doc, "Jordan Elliot Whitfield (DOB 12/04/1985, Medicare 2934 5671 0)", bold_label="Client: ")
    doc.add_paragraph()

    _heading(doc, "Attendees", level=2)
    _grid_table(doc, ["Name", "Role", "Organisation"], [
        ["Jordan Whitfield", "Client", "—"],
        ["Priya Whitfield", "Support person (spouse)", "—"],
        ["Dr. Amelia Ferro", "Clinical Psychologist (treating)", "Northgate Psychology Clinic"],
        ["Dr. Susan Ng", "General Practitioner", "Riverside Family Medical Practice"],
        ["Grace Tan", "Occupational Therapist", "ReturnAbility OT Services"],
        ["Marcus Delaney", "Rehabilitation Case Manager", "VicWork Insurance"],
        ["Helen Okafor", "Return to Work Coordinator", "Coburg Logistics Pty Ltd"],
        ["Adrian Cross", "Team Leader, Warehouse Operations", "Coburg Logistics Pty Ltd"],
    ])
    doc.add_paragraph()

    _heading(doc, "Background", level=2)
    doc.add_paragraph(
        "Convened to review Mr Whitfield's progress under his current Mental Health Care Plan and "
        "WorkCover claim following the workplace incident of 14 October 2025, and to agree a "
        "coordinated plan for graduated return to full warehouse-floor duties. Mr Whitfield has "
        "completed 5 of 6 funded psychology sessions of trauma-focused CBT with graded exposure, "
        "with standardised measures (DASS-21, PCL-5) showing consistent improvement, and has "
        "completed step 1 of a 6-step exposure hierarchy for the warehouse floor."
    )

    _heading(doc, "Discussion", level=2)
    doc.add_paragraph(
        "Dr Ferro summarised treatment progress, noting reduced intrusive memories (from "
        "3–4/week to approximately 1/week) and a PCL-5 score now below the clinical cutoff. She "
        "advised that continued graded exposure, coordinated with a supervised return to the "
        "warehouse floor, offers the best prospect of full duties resuming within 2–3 months. "
        "Ms Tan (OT) outlined a functional capacity assessment supporting a graded increase in "
        "floor-based hours, beginning with short supervised periods alongside a colleague. "
        "Mr Delaney confirmed VicWork Insurance's support for extending funded psychology sessions "
        "beyond the current Mental Health Care Plan allocation, pending a further GP review. "
        "Ms Okafor and Mr Cross confirmed Coburg Logistics can continue modified/office-based "
        "duties in the interim and will roster a familiar colleague to accompany Mr Whitfield "
        "during initial floor-exposure sessions. Mr Whitfield reported feeling supported by the "
        "plan and reiterated his goal of returning to full supervisory duties before the end of "
        "the current financial year."
    )

    _heading(doc, "Decisions / Actions", level=2)
    doc.add_paragraph(
        "It was agreed that Dr Susan Ng would provide a further review letter to enable extended "
        "Mental Health Care Plan sessions within 1 week, and that Dr Amelia Ferro would continue "
        "trauma-focused CBT with graded exposure through steps 2–6 on an ongoing weekly basis. "
        "Grace Tan will complete a functional capacity review ahead of any increase in floor-duty "
        "hours within 2 weeks, and Marcus Delaney will approve the extension of funded psychology "
        "sessions under the claim within 5 business days. Helen Okafor and Adrian Cross will roster "
        "supervised floor-exposure sessions with a familiar colleague from the next roster cycle. "
        "A follow-up case conference will be scheduled by Dr Amelia Ferro in 8 weeks (mid-May 2026) "
        "to review progress against these actions."
    )
    doc.add_paragraph()

    _heading(doc, "Next Case Conference", level=2)
    doc.add_paragraph(
        "Tentatively scheduled for mid-May 2026, or earlier if a change in risk or function is "
        "identified by any party. Minutes to be circulated to all attendees and filed in the "
        "client's record at Northgate Psychology Clinic and Riverside Family Medical Practice."
    )

    doc.save(path)


if __name__ == "__main__":
    build_referral_letter(os.path.join(OUT_DIR, "01_gp_referral_letter.docx"))
    build_intake_notes(os.path.join(OUT_DIR, "02_biopsychosocial_intake_notes.docx"))
    build_session_log(os.path.join(OUT_DIR, "03_session_log_progress_notes.docx"))
    build_treatment_review_source(os.path.join(OUT_DIR, "04_treatment_review_source.docx"))
    build_discharge_summary(os.path.join(OUT_DIR, "05_discharge_summary.docx"))
    build_risk_assessment(os.path.join(OUT_DIR, "06_risk_assessment.docx"))
    build_case_conference_note(os.path.join(OUT_DIR, "07_case_conference_note.docx"))
    print("done:", os.listdir(OUT_DIR))
