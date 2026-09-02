"""Infer a :class:`FormSpec` from an arbitrary table-based ``.docx`` form, so
a clinic can add its own template instead of only the three bundled ones.

The row-classification and grid logic already live in ``clinical_forms``
(:func:`_walk_table`, :func:`_grid_fields`) and are reused verbatim. This
module only supplies the three things the bundled specs hand-code:

* the header block — which top rows are ``label:`` metadata cells, and where
  their values land;
* the body start row — the first row after the header block;
* the grid ranges — a banner + column-header row + row-header column.

An uploaded template is stored as its own ``.docx`` under
``<app_data>/templates/<form_id>.docx`` (the file is the source of truth,
re-parsed on load) with a tiny ``<form_id>.json`` sidecar for the selector.
Nothing here calls a model or the network.
"""

from __future__ import annotations

import io
import json
from pathlib import Path

import docx

from . import clinical_forms as cf
from . import desktop

# slug (from clinical_forms.slugify) -> (canonical key, display label)
_KNOWN_HEADERS: dict[str, tuple[str, str]] = {
    "date": ("date", "Date"),
    "practitioner": ("practitioner", "Practitioner"),
    "client_name": ("client_name", "Client name"),
    "client_dob": ("client_dob", "Client DOB"),
    "session_number": ("session_number", "Session number"),
    "session_no": ("session_number", "Session number"),
    "item_code": ("item_code", "Item code (if relevant)"),
}


def _match_header(text: str) -> tuple[str, str] | None:
    slug = cf.slugify(text)
    for prefix, canonical in _KNOWN_HEADERS.items():
        if slug == prefix or slug.startswith(prefix + "_"):
            return canonical
    return None


def _is_blank_row(cells) -> bool:
    return all(not c.text.strip() for c in cells)


def _skip_blank(table, i: int) -> int:
    n = len(table.rows)
    while i < n and _is_blank_row(cf._dedupe_row(table.rows[i])):
        i += 1
    return i


def _infer_header(table) -> tuple[list, int]:
    """Return ``(header_fields, body_start_row)`` for the first table.

    Header rows are consecutive ``label:`` metadata rows whose every non-blank
    cell matches a known identifier label, optionally followed by a full-width
    ``Reason for referral:`` label row (its answer lives in the blank row
    beneath it). The body starts at the first row that is neither.
    """
    n = len(table.rows)
    headers: list = []
    i = _skip_blank(table, 0)

    while i < n:
        cells = cf._dedupe_row(table.rows[i])
        non_blank = [(c, ci) for ci, c in enumerate(cells) if c.text.strip()]
        matches = [(_match_header(c.text), ci) for c, ci in non_blank]
        if non_blank and all(m for m, _ in matches):
            for (key, label), ci in ((m, ci) for m, ci in matches):
                headers.append(cf.HeaderField(key, label, 0, i, ci, "inline"))
            i += 1
            # a single blank row between two metadata rows is still header
            look = _skip_blank(table, i)
            if look > i and look < n:
                nxt = cf._dedupe_row(table.rows[look])
                nb = [c for c in nxt if c.text.strip()]
                if nb and all(_match_header(c.text) for c in nb):
                    i = look
            continue
        break

    body_start = i
    j = _skip_blank(table, i)
    if j < n:
        cells = cf._dedupe_row(table.rows[j])
        if len(cells) == 1 and cf.slugify(cells[0].text).startswith("reason_for"):
            headers.append(
                cf.HeaderField("reason_for_referral", "Reason for referral", 0, j, 0, "append")
            )
            body_start = _skip_blank(table, j + 1)

    return headers, body_start


def _short(text: str) -> bool:
    return 0 < len(text.split()) <= 4


def _find_grids(table, start_row: int) -> list[tuple[int, int, int, int, str]]:
    """Locate ``(banner_row, header_row, first_data_row, last_data_row, section)``
    for every row-label x column-label grid at or after ``start_row``.

    Shape: a full-width banner, then a row whose first cell is empty and whose
    remaining cells are short column headers, then >= 2 rows whose first cell is
    a label and whose remaining cells are blank.
    """
    grids: list[tuple[int, int, int, int, str]] = []
    n = len(table.rows)
    i = start_row
    while i < n - 2:
        banner = cf._dedupe_row(table.rows[i])
        header = cf._dedupe_row(table.rows[i + 1])
        is_banner = len(banner) == 1 and banner[0].text.strip()
        is_grid_header = (
            len(header) >= 3
            and not header[0].text.strip()
            and all(c.text.strip() and _short(c.text) for c in header[1:])
        )
        if is_banner and is_grid_header:
            k = i + 2
            while k < n:
                row = cf._dedupe_row(table.rows[k])
                if len(row) >= 2 and row[0].text.strip() and all(not c.text.strip() for c in row[1:]):
                    k += 1
                else:
                    break
            if k - (i + 2) >= 2:
                grids.append((i, i + 1, i + 2, k - 1, banner[0].text.strip()))
                i = k
                continue
        i += 1
    return grids


def _parse_document(doc, *, form_id: str, title: str, asset_path: Path) -> cf.FormSpec:
    if not doc.tables:
        raise cf.ClinicalFormError(
            "This .docx has no table. CareScribe fills table-based forms only."
        )
    header_fields, body_start = _infer_header(doc.tables[0])
    fields: list = []
    for table_index, table in enumerate(doc.tables):
        cursor = body_start if table_index == 0 else 0
        for banner_row, header_row, first_data, last_data, section in _find_grids(table, cursor):
            if banner_row > cursor:
                fields += cf._walk_table(table, table_index, cursor, end_row=banner_row)
            fields += cf._grid_fields(
                table, table_index, header_row, first_data, last_data, section
            )
            cursor = last_data + 1
        fields += cf._walk_table(table, table_index, cursor)

    if not fields:
        raise cf.ClinicalFormError(
            "No fillable fields were found in this template. Expected a table of "
            "label cells with blank cells beside them."
        )

    _reject_duplicate_keys(fields, form_id)
    return cf.FormSpec(
        form_id=form_id, title=title, asset_path=asset_path,
        header_fields=header_fields, fields=fields,
    )


def _reject_duplicate_keys(fields, form_id: str) -> None:
    seen: set[str] = set()
    for f in fields:
        if f.key in seen:
            raise cf.ClinicalFormError(
                f"Template '{form_id}' produced two fields with the same key "
                f"('{f.key}'). Rename one of the repeated section or row labels."
            )
        seen.add(f.key)


def _title_from(doc, fallback_name: str) -> str:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            return text
    return Path(fallback_name).stem or "Custom template"


# --- public parsing entry points ------------------------------------------------

def parse_template_bytes(data: bytes, *, form_id: str, title: str | None = None) -> cf.FormSpec:
    doc = docx.Document(io.BytesIO(data))
    return _parse_document(
        doc, form_id=form_id, title=title or _title_from(doc, form_id),
        asset_path=_templates_dir() / f"{form_id}.docx",
    )


def parse_template_path(path, *, form_id: str, title: str | None = None) -> cf.FormSpec:
    path = Path(path)
    doc = docx.Document(path)
    return _parse_document(
        doc, form_id=form_id, title=title or _title_from(doc, path.name), asset_path=path,
    )


def fill_parsed_template(
    spec: cf.FormSpec, original_docx: bytes,
    field_values: dict[str, str], header_values: dict[str, str],
) -> bytes:
    """Fill an in-memory copy of ``original_docx`` from ``spec``'s anchors.

    A thin wrapper over the write logic in ``clinical_forms.fill_template`` for
    the case where the template bytes are in hand rather than on disk (upload
    preview). Nothing touches disk.
    """
    from dataclasses import replace

    doc_spec = replace(spec, asset_path=io.BytesIO(original_docx))
    return cf.fill_template(doc_spec, field_values, header_values)


# --- persistence + registry ---------------------------------------------------

def _templates_dir() -> Path:
    return desktop.app_data_dir() / "templates"


def _slug_id(name: str) -> str:
    return cf.slugify(Path(name).stem) or "custom_template"


def _unique_id(base: str) -> str:
    taken = set(cf._FORM_SPEC_BUILDERS)
    directory = _templates_dir()
    if directory.exists():
        taken |= {p.stem for p in directory.glob("*.docx")}
    if base not in taken:
        return base
    n = 2
    while f"{base}_{n}" in taken:
        n += 1
    return f"{base}_{n}"


def save_template(data: bytes, original_name: str) -> str:
    """Validate an uploaded ``.docx``, store it, and return its new form id.

    Raises :class:`ClinicalFormError` if it does not parse into a usable form.
    """
    doc = docx.Document(io.BytesIO(data))
    title = _title_from(doc, original_name)
    form_id = _unique_id(_slug_id(original_name))
    directory = _templates_dir()
    asset_path = directory / f"{form_id}.docx"

    # Fail before writing anything if the template is unusable.
    _parse_document(doc, form_id=form_id, title=title, asset_path=asset_path)

    directory.mkdir(parents=True, exist_ok=True)
    asset_path.write_bytes(data)
    (directory / f"{form_id}.json").write_text(
        json.dumps({"form_id": form_id, "title": title}), encoding="utf-8"
    )
    cf.get_form_spec.cache_clear()
    return form_id


def delete_template(form_id: str) -> None:
    directory = _templates_dir()
    for suffix in (".docx", ".json"):
        (directory / f"{form_id}{suffix}").unlink(missing_ok=True)
    cf.get_form_spec.cache_clear()


def user_form_options() -> list[tuple[str, str]]:
    directory = _templates_dir()
    if not directory.exists():
        return []
    options: list[tuple[str, str]] = []
    for sidecar in sorted(directory.glob("*.json")):
        try:
            meta = json.loads(sidecar.read_text(encoding="utf-8"))
            if (directory / f"{meta['form_id']}.docx").is_file():
                options.append((meta["form_id"], meta["title"]))
        except (OSError, ValueError, KeyError):
            continue
    return options


def load_user_spec(form_id: str) -> cf.FormSpec | None:
    asset_path = _templates_dir() / f"{form_id}.docx"
    if not asset_path.is_file():
        return None
    return parse_template_path(asset_path, form_id=form_id)
