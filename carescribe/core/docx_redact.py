"""
Structure-preserving .docx redaction.

apply_redactions(path_in, path_out, replacements) takes an APPROVED mapping of
{literal_string: placeholder} and rewrites the document so every occurrence of each
literal (including ones split across Word's fragmented runs, and names wrapped across a
line/paragraph break) becomes its placeholder — while leaving all tables, styles,
headings, headers and footers structurally untouched.

Design notes:
- Word splits one visible word across many <w:r> runs. We reconstruct each paragraph's
  text from its runs, match on that, then write the placeholder into the first run of the
  span and blank the overlapped tail of later runs. Formatting of the first run is kept.
- Replacements are applied longest-literal-first and right-to-left so offsets never shift
  and a short literal can't eat part of a longer one.
- Cross-paragraph fallback handles a name wrapped onto the next paragraph (P_i suffix +
  P_{i+1} prefix == literal, ignoring whitespace).
- This module does NOT decide what is an identifier. It only applies the human-approved
  map. Detection stays in the de-id pipeline; this keeps the Word output identical to what
  the reviewer approved.

Vendored as provided, unchanged apart from :func:`has_unreachable_text`, which the UI
needs in order to warn about text boxes rather than pass them silently.
"""
from docx import Document


def _iter_groups(doc):
    """Yield lists of paragraphs, each list being one 'flow' where wrapping can occur
    (body, a table cell, a header/footer). Cross-paragraph stitching stays inside a group
    so unrelated paragraphs (e.g. across a table boundary) are never joined."""
    def walk(container):
        paras = list(getattr(container, "paragraphs", []))
        if paras:
            yield paras
        for t in getattr(container, "tables", []):
            for row in t.rows:
                for cell in row.cells:
                    yield from walk(cell)
    yield from walk(doc)
    for section in doc.sections:
        for hf in (section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer):
            if hf is not None:
                yield from walk(hf)


def _iter_paragraphs(doc):
    for group in _iter_groups(doc):
        yield from group


def _replace_in_paragraph(paragraph, literals):
    """Apply literals (list of (text, placeholder), longest-first) within one paragraph."""
    runs = paragraph.runs
    if not runs:
        return
    # char index -> run index map over the concatenated paragraph text
    text = ""
    owner = []            # owner[i] = index of the run that contributed char i
    for ri, r in enumerate(runs):
        rt = r.text or ""
        text += rt
        owner.extend([ri] * len(rt))
    if not text.strip():
        return

    # collect non-overlapping matches, scanning longest literals first
    spans = []            # (start, end, placeholder)
    claimed = [False] * len(text)
    for lit, ph in literals:
        if not lit:
            continue
        start = 0
        while True:
            idx = text.find(lit, start)
            if idx == -1:
                break
            end = idx + len(lit)
            if not any(claimed[idx:end]):
                spans.append((idx, end, ph))
                for k in range(idx, end):
                    claimed[k] = True
            start = idx + 1
    if not spans:
        return

    # rebuild each run's text, right-to-left so indices stay valid
    run_text = [r.text or "" for r in runs]
    for start, end, ph in sorted(spans, key=lambda s: s[0], reverse=True):
        first_run = owner[start]
        # offset of `start` within its run
        run_start_char = start - next(i for i in range(start, -1, -1)
                                      if i == 0 or owner[i - 1] != first_run)
        # clear the span across all runs it touches
        for ci in range(start, end):
            ri = owner[ci]
            local = ci - next(j for j in range(ci, -1, -1)
                              if j == 0 or owner[j - 1] != ri)
            s = run_text[ri]
            run_text[ri] = s[:local] + "\x00" + s[local + 1:]
        # drop the marked chars, then inject the placeholder at the span's first run
        for ri in {owner[c] for c in range(start, end)}:
            run_text[ri] = run_text[ri].replace("\x00", "")
        rs = run_text[first_run]
        run_text[first_run] = rs[:run_start_char] + ph + rs[run_start_char:]

    for r, nt in zip(runs, run_text):
        if r.text != nt:
            r.text = nt


def _norm(s):
    return " ".join(s.split())


def _replace_across_paragraphs(paragraphs, literals):
    """Redact a literal split across a paragraph boundary (wrapped name)."""
    texts = [p.text for p in paragraphs]
    for i in range(len(paragraphs) - 1):
        a, b = texts[i], texts[i + 1]
        if not a or not b:
            continue
        for lit, ph in literals:
            n = _norm(lit)
            # try: some suffix of A + some prefix of B, whitespace-insensitive, equals lit
            for cut_a in range(len(a)):
                sa = a[cut_a:]
                if not sa.strip():
                    continue
                if not _norm(lit).startswith(_norm(sa)):
                    continue
                remaining = _norm(lit)[len(_norm(sa)):].strip()
                if remaining and _norm(b).startswith(remaining):
                    # place full placeholder in A's tail, delete matched prefix in B
                    _replace_in_paragraph(paragraphs[i], [(sa.strip(), ph)])
                    # delete the consumed head of B (first `len(remaining)` non-space chars)
                    _delete_prefix(paragraphs[i + 1], remaining)
                    texts[i] = paragraphs[i].text
                    texts[i + 1] = paragraphs[i + 1].text
                    break


def _delete_prefix(paragraph, normalized_prefix):
    """Delete the leading text of a paragraph matching normalized_prefix (ws-insensitive)."""
    want = _norm(normalized_prefix)
    if not want:
        return
    acc = ""
    for r in paragraph.runs:
        rt = r.text or ""
        new_chars = []
        for ch in rt:
            if _norm(acc) == want:
                new_chars.append(ch)
            elif ch.isspace():
                new_chars.append("" if _norm(acc) != want else ch)
                acc += ch
            else:
                if len(_norm(acc + ch)) <= len(want):
                    acc += ch  # consume (delete)
                else:
                    new_chars.append(ch)
        r.text = "".join(new_chars)
        if _norm(acc) == want:
            # leave the rest of the runs untouched
            continue


def apply_redactions(path_in, path_out, replacements):
    """replacements: dict {literal_string: placeholder}. Writes redacted docx to path_out."""
    doc = Document(path_in)
    # longest first so a longer identifier is matched before any substring of it
    literals = sorted(replacements.items(), key=lambda kv: len(kv[0]), reverse=True)

    groups = list(_iter_groups(doc))
    # 1) stitch names wrapped across a paragraph boundary FIRST, within each flow,
    #    so a standalone first-name literal can't consume half of a wrapped full name.
    for group in groups:
        _replace_across_paragraphs(group, literals)
    # 2) then the per-paragraph, multi-run pass for everything else.
    for group in groups:
        for p in group:
            _replace_in_paragraph(p, literals)
    doc.save(path_out)
    return path_out


def has_unreachable_text(path) -> bool:
    """True if the document holds text this module cannot reach.

    Text boxes, shapes/SmartArt and embedded objects live outside the paragraph tree
    walked above, so their contents survive redaction untouched. Callers must warn a
    reviewer rather than let such a file pass silently. Added for the UI's text-box
    warning; the redaction logic above is the vendored original.
    """
    doc = Document(path)
    xml = doc.element.body.xml
    for section in doc.sections:
        for hf in (section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer):
            if hf is not None:
                xml += hf._element.xml
    return "txbxContent" in xml or "<w:drawing" in xml or "<w:object" in xml


def extract_text(path):
    """Flatten a docx to text (body + tables + headers/footers) for a residual scan."""
    doc = Document(path)
    return "\n".join(p.text for p in _iter_paragraphs(doc))
