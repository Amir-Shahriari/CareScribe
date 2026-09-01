"""
Fill the three bundled APS clinical form templates from approved,
de-identified source documents.

Each template's fillable cells are discovered by walking its real table
structure (see docs/superpowers/specs/2026-08-13-clinical-forms-design.md),
not hand-transcribed, because merged cells make column indices unreliable
by eye. Generation asks the model for one delimited block of text per
discovered field; nothing here decides what counts as a field at
generation time — that was fixed when the spec was built from the asset.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path

import docx

from . import mapping as _mapping  # PLACEHOLDER_RE, reused rather than reimplemented

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"


class ClinicalFormError(RuntimeError):
    """Raised when a clinical form can't be built or filled."""


def slugify(text: str) -> str:
    text = (text or "").lower().strip()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


@dataclass(frozen=True)
class FormField:
    key: str
    label: str
    table_index: int
    value_row_index: int
    value_col_index: int
    append_after_label: bool


@dataclass(frozen=True)
class HeaderField:
    key: str
    label: str
    table_index: int
    row_index: int
    col_index: int
    style: str  # "inline" (append to the label's own paragraph) or "append"
                # (new paragraph(s) after the label, same cell)


@dataclass(frozen=True)
class FormSpec:
    form_id: str
    title: str
    asset_path: Path
    header_fields: list[HeaderField]
    fields: list[FormField]


def _dedupe_row(row):
    """Deduplicate a row's cells by underlying XML element identity.

    python-docx reports a merged cell once per grid column it spans, so a
    naive ``row.cells`` walk sees the same cell 2-3 times. Comparing
    ``id(cell._tc)`` collapses those back into one logical cell per visual
    box, in left-to-right order.
    """
    seen: set[int] = set()
    cells = []
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            cells.append(cell)
    return cells


def _paragraph_texts(cell) -> list[str]:
    return [p.text for p in cell.paragraphs]


def _walk_table(
    table, table_index: int, start_row: int, header_seed: str = "",
    end_row: int | None = None,
) -> list[FormField]:
    fields: list[FormField] = []
    current_header = header_seed
    pending: tuple[int, str] | None = None  # (row_index, label_text)

    def make_key(label: str) -> str:
        return f"{slugify(current_header)}.{slugify(label)}" if current_header else slugify(label)

    stop = len(table.rows) if end_row is None else min(end_row, len(table.rows))
    for row_index in range(start_row, stop):
        cells = _dedupe_row(table.rows[row_index])
        texts = [c.text.strip() for c in cells]
        full_blank = all(t == "" for t in texts)

        if pending is not None:
            _, label_text = pending
            pending = None
            if full_blank and len(cells) == 1:
                fields.append(FormField(
                    key=make_key(label_text), label=label_text,
                    table_index=table_index, value_row_index=row_index,
                    value_col_index=0, append_after_label=False,
                ))
                continue
            current_header = label_text
            # fall through — this row still needs its own classification

        if full_blank:
            continue

        if len(cells) == 1:
            paragraphs = _paragraph_texts(cells[0])
            label_text = paragraphs[0].strip()
            trailing = paragraphs[1:]
            trailing_blank = bool(trailing) and all(p.strip() == "" for p in trailing)
            if trailing_blank:
                fields.append(FormField(
                    key=make_key(label_text), label=label_text,
                    table_index=table_index, value_row_index=row_index,
                    value_col_index=0, append_after_label=True,
                ))
            elif len(paragraphs) == 1:
                pending = (row_index, label_text)
            else:
                current_header = label_text
            continue

        label_text = cells[0].text.strip()
        if label_text.lower().startswith("signature"):
            continue
        fields.append(FormField(
            key=make_key(label_text), label=label_text,
            table_index=table_index, value_row_index=row_index,
            value_col_index=1, append_after_label=False,
        ))

    return fields


def _treatment_review_spec() -> FormSpec:
    asset = TEMPLATES_DIR / "client_treatment_review.docx"
    doc = docx.Document(asset)
    (table,) = doc.tables
    fields = _walk_table(table, table_index=0, start_row=6)
    header_fields = [
        HeaderField("date", "Date", 0, 0, 0, "inline"),
        HeaderField("practitioner", "Practitioner", 0, 0, 1, "inline"),
        HeaderField("client_name", "Client name", 0, 1, 0, "inline"),
        HeaderField("client_dob", "Client DOB", 0, 1, 1, "inline"),
        HeaderField("session_number", "Session number", 0, 2, 0, "inline"),
        HeaderField("item_code", "Item code (if relevant)", 0, 2, 1, "inline"),
        HeaderField("reason_for_referral", "Reason for referral", 0, 4, 0, "append"),
    ]
    return FormSpec(
        form_id="client_treatment_review", title="Client Treatment Review",
        asset_path=asset, header_fields=header_fields, fields=fields,
    )


def _grid_fields(table, table_index: int, header_row: int, first_data_row: int,
                  last_data_row: int, section_header: str) -> list[FormField]:
    """The Biopsychosocial 'CLINICAL FORMULATION' table: a row-label ×
    column-label grid, not a single label/value pair per row."""
    header_cells = _dedupe_row(table.rows[header_row])
    column_labels = [c.text.strip() for c in header_cells[1:]]
    fields = []
    for row_index in range(first_data_row, last_data_row + 1):
        cells = _dedupe_row(table.rows[row_index])
        row_label = cells[0].text.strip()
        for col_offset, column_label in enumerate(column_labels, start=1):
            key = f"{slugify(section_header)}.{slugify(row_label)}.{slugify(column_label)}"
            fields.append(FormField(
                key=key, label=f"{row_label} – {column_label}",
                table_index=table_index, value_row_index=row_index,
                value_col_index=col_offset, append_after_label=False,
            ))
    return fields


def _session_notes_spec() -> FormSpec:
    asset = TEMPLATES_DIR / "client_session_notes.docx"
    doc = docx.Document(asset)
    (table,) = doc.tables
    fields = _walk_table(table, table_index=0, start_row=6)
    header_fields = [
        HeaderField("date", "Date", 0, 0, 0, "inline"),
        HeaderField("practitioner", "Practitioner", 0, 0, 1, "inline"),
        HeaderField("client_name", "Client name", 0, 1, 0, "inline"),
        HeaderField("client_dob", "Client DOB", 0, 1, 1, "inline"),
        HeaderField("session_number", "Session Number", 0, 2, 0, "inline"),
        HeaderField("item_code", "Item code (if relevant)", 0, 2, 1, "inline"),
        HeaderField("reason_for_referral", "Reason for referral", 0, 4, 0, "append"),
    ]
    return FormSpec(
        form_id="client_session_notes", title="Client Session Notes",
        asset_path=asset, header_fields=header_fields, fields=fields,
    )


def _biopsychosocial_spec() -> FormSpec:
    asset = TEMPLATES_DIR / "biopsychosocial_assessment.docx"
    doc = docx.Document(asset)
    table0, table1 = doc.tables
    fields = _walk_table(table0, table_index=0, start_row=5)
    fields += _grid_fields(
        table1, table_index=1, header_row=1, first_data_row=2, last_data_row=5,
        section_header="CLINICAL FORMULATION",
    )
    fields += _walk_table(table1, table_index=1, start_row=6)
    header_fields = [
        HeaderField("date", "Date", 0, 0, 0, "inline"),
        HeaderField("practitioner", "Practitioner", 0, 0, 1, "inline"),
        HeaderField("client_name", "Client Name", 0, 1, 0, "inline"),
        HeaderField("client_dob", "Client DOB", 0, 1, 1, "inline"),
        HeaderField("reason_for_referral", "Reason for referral", 0, 3, 0, "append"),
    ]
    return FormSpec(
        form_id="biopsychosocial_assessment", title="Biopsychosocial Assessment",
        asset_path=asset, header_fields=header_fields, fields=fields,
    )


_FORM_SPEC_BUILDERS = {
    "client_treatment_review": _treatment_review_spec,
}

_FORM_SPEC_BUILDERS["client_session_notes"] = _session_notes_spec
_FORM_SPEC_BUILDERS["biopsychosocial_assessment"] = _biopsychosocial_spec


@lru_cache(maxsize=None)
def get_form_spec(form_id: str) -> FormSpec:
    builder = _FORM_SPEC_BUILDERS.get(form_id)
    if builder is not None:
        return builder()
    from . import template_ingest  # lazy — template_ingest imports this module

    spec = template_ingest.load_user_spec(form_id)
    if spec is None:
        raise ClinicalFormError(f"Unknown clinical form '{form_id}'.")
    return spec


def available_forms() -> list[tuple[str, str]]:
    """(form_id, title) pairs — bundled forms first, then clinic-uploaded ones."""
    forms = [(form_id, builder().title) for form_id, builder in _FORM_SPEC_BUILDERS.items()]
    try:
        from . import template_ingest

        forms += template_ingest.user_form_options()
    except Exception:  # noqa: BLE001 — a broken user template must not hide the bundled forms
        pass
    return forms


def _clear_cell(cell) -> None:
    """Remove every paragraph after the first, and every run in the first,
    leaving one empty paragraph ready to receive fresh text."""
    paragraphs = cell.paragraphs
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for run in list(paragraphs[0].runs):
        run._element.getparent().remove(run._element)


def _fill_cell(cell, text: str) -> None:
    """Overwrite a dedicated value cell (label lives in a different cell)."""
    _clear_cell(cell)
    text = (text or "").strip() or "Not documented"
    lines = text.splitlines() or ["Not documented"]
    cell.paragraphs[0].add_run(lines[0])
    for line in lines[1:]:
        cell.add_paragraph(line)


def _fill_cell_after_label(cell, text: str) -> None:
    """Append text as new paragraphs after an existing label paragraph,
    which must survive untouched (used for own-cell fields and the
    'Reason for referral' header field)."""
    paragraphs = cell.paragraphs
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    text = (text or "").strip() or "Not documented"
    for line in text.splitlines() or ["Not documented"]:
        cell.add_paragraph(line)


def _fill_header_cell(cell, value: str) -> None:
    """Append a typed value inline, on the label's own paragraph and run —
    'Date: ' becomes 'Date: 12/08/2026'."""
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[-1].text = paragraph.runs[-1].text + value
    else:
        paragraph.add_run(value)


def fill_template(
    form_spec: FormSpec, field_values: dict[str, str], header_values: dict[str, str]
) -> bytes:
    """Fill a fresh in-memory copy of the template. Nothing touches disk."""
    doc = docx.Document(form_spec.asset_path)

    for header in form_spec.header_fields:
        value = (header_values.get(header.key) or "").strip()
        if not value:
            continue
        try:
            cell = _dedupe_row(
                doc.tables[header.table_index].rows[header.row_index]
            )[header.col_index]
        except IndexError:
            raise ClinicalFormError(
                f"Template shape mismatch for form '{form_spec.form_id}': header "
                f"'{header.key}' expects table {header.table_index}, row "
                f"{header.row_index}, column {header.col_index}, which does not "
                "exist in this template."
            ) from None
        if header.style == "inline":
            _fill_header_cell(cell, value)
        else:
            _fill_cell_after_label(cell, value)

    for field in form_spec.fields:
        try:
            cell = _dedupe_row(
                doc.tables[field.table_index].rows[field.value_row_index]
            )[field.value_col_index]
        except IndexError:
            raise ClinicalFormError(
                f"Template shape mismatch for form '{form_spec.form_id}': field "
                f"'{field.key}' expects table {field.table_index}, row "
                f"{field.value_row_index}, column {field.value_col_index}, which "
                "does not exist in this template."
            ) from None
        text = (field_values.get(field.key) or "").strip() or "Not documented"
        if field.append_after_label:
            _fill_cell_after_label(cell, text)
        else:
            _fill_cell(cell, text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


_SYSTEM_PREAMBLE = """\
You are a clinical documentation assistant filling out a structured clinical \
form from de-identified source material. You produce a DRAFT that a \
clinician will review, correct, and sign.

The source text has already been de-identified. Names, numbers, dates and \
places have been replaced with bracketed placeholders such as [PATIENT], \
[CLINICIAN_1], [MRN] and [DATE_2].

Rules you must follow:

1. Use ONLY information present in the provided source text.
2. Do NOT invent facts, names, dates, dosages, values, or events.
3. Preserve every bracketed placeholder EXACTLY as written, character for \
character. Do not translate, expand, renumber, or replace them.
4. Write in plain clinical register, past tense, concisely.
5. For each field listed below, write the field's marker on its own line, \
then the field's content, then move to the next field. If the source has \
no information for a field, write exactly "Not documented" for that field \
rather than guessing or leaving it blank.
6. Do not add commentary, a greeting, or a sign-off. Output only the fields.
7. Some fields carry one or more "house-style example" lines. Match their \
tone, length, and structure. They are from other de-identified notes — take \
STYLE from them only, never facts. Every clinical fact must come from the \
source text above.

FIELDS TO COMPLETE, IN THIS EXACT ORDER — reproduce each marker exactly:

{field_list}
"""

_USER_TEMPLATE = """\
SOURCE TEXT (de-identified):
---
{document}
---

Write the content for every field listed in the system instructions, each \
under its own <<FIELD:key>> marker, in the exact order given.
"""


def build_prompt(
    form_spec: FormSpec,
    deidentified_text: str,
    exemplars: dict[str, list[str]] | None = None,
) -> tuple[str, str]:
    """Build the (system, user) prompt pair.

    ``exemplars`` maps a field key to house-style example values (de-identified
    text from this clinic's own past drafts); each is rendered as an indented
    line under its field so the model matches style without borrowing facts.
    """
    exemplars = exemplars or {}
    lines: list[str] = []
    for field in form_spec.fields:
        lines.append(f"<<FIELD:{field.key}>> — {field.label}")
        for example in exemplars.get(field.key, []):
            flat = " ".join(str(example).split())
            lines.append(f"    house-style example: {flat}")
    system = _SYSTEM_PREAMBLE.format(field_list="\n".join(lines))
    user = _USER_TEMPLATE.format(document=deidentified_text)
    return system, user


_FIELD_MARKER_RE = re.compile(r"<<FIELD:([a-z0-9_.]+)>>")


def parse_fields(form_spec: FormSpec, raw_output: str) -> dict[str, str]:
    """Turn the model's marker-delimited output into ``{field_key: text}``.

    Any field the model skipped defaults to "Not documented" — enforced
    here rather than trusted to the model. Unknown markers are ignored; on
    a duplicate marker for the same key, the first occurrence wins.
    """
    text = raw_output or ""
    matches = list(_FIELD_MARKER_RE.finditer(text))
    found: dict[str, str] = {}
    for index, match in enumerate(matches):
        key = match.group(1)
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        if key not in found:
            found[key] = text[start:end].strip().strip("—-").strip()

    return {
        field.key: found.get(field.key) or "Not documented"
        for field in form_spec.fields
    }


def combine_sources(
    sources: list[tuple[str, str, dict[str, str]]]
) -> tuple[str, dict[str, str]]:
    """Concatenate several documents' de-identified text into one source.

    ``sources`` is ``(name, redacted_text, phi_map)`` per contributing
    document. Every placeholder is prefixed with that document's position
    (``[PATIENT]`` -> ``[DOCA_PATIENT]``) so two documents that each define
    their own ``[PATIENT]`` never collide once merged — each keeps
    resolving to its own real value at re-identification time.

    Raises ``ClinicalFormError`` if sources is empty or exceeds 26 documents.
    """
    if not sources:
        raise ClinicalFormError("No source documents were selected.")
    if len(sources) > 26:
        raise ClinicalFormError("Maximum 26 source documents permitted.")

    parts = []
    merged: dict[str, str] = {}
    for index, (name, text, phi_map) in enumerate(sources, start=1):
        # Use letter-based tag (DOCA_, DOCB_, ...) to preserve trailing digit
        # run position expected by mapping.PLACEHOLDER_RE regex pattern.
        prefix = f"DOC{chr(ord('A') + index - 1)}_"
        rewritten = text or ""
        # Drive both text and map rewrite from phi_map.keys() to ensure
        # consistency: every placeholder in map gets rewritten in text,
        # and vice versa, regardless of whether it matches PLACEHOLDER_RE.
        for placeholder, value in (phi_map or {}).items():
            new_key = f"[{prefix}{placeholder[1:]}"
            rewritten = rewritten.replace(placeholder, new_key)
            merged[new_key] = value

        # Use plain positional separator; do not interpolate filename
        # which may contain PII from the original upload.
        parts.append(f"--- Source {index} ---\n{rewritten}")

    return "\n\n".join(parts), merged


from typing import Iterable, Iterator

from . import carenotes


_PLACEHOLDER_TOKEN_RE = re.compile(r"\[[A-Z][A-Z0-9_]*\]")


def _form_grammar(form_spec: FormSpec, combined_text: str) -> str | None:
    """A GBNF grammar that requires every ``<<FIELD:key>>`` marker in order and
    forbids a bracket token that is not one of this document's placeholders."""
    from . import grammar as _grammar

    placeholders = sorted(set(_PLACEHOLDER_TOKEN_RE.findall(combined_text or "")))
    return _grammar.field_grammar([f.key for f in form_spec.fields], placeholders)


def generate_form_document(
    combined_text: str,
    form_spec: FormSpec,
    backend,
    stream: bool = True,
    *,
    phi_values: Iterable[str] | None = None,
    acknowledged: Iterable[str] = (),
    exemplars: dict[str, list[str]] | None = None,
) -> Iterator[str]:
    system, user = build_prompt(form_spec, combined_text, exemplars)
    return carenotes.generate_document(
        combined_text, form_spec.form_id, backend, stream,
        phi_values=phi_values, acknowledged=acknowledged,
        system=system, user_prompt=user,
        grammar=_form_grammar(form_spec, combined_text),
    )


def refine_form_document(
    combined_text: str,
    draft_marker_text: str,
    instruction: str,
    form_spec: FormSpec,
    backend,
    stream: bool = True,
    *,
    history: list[tuple[str, str]] | None = None,
    phi_values: Iterable[str] | None = None,
    acknowledged: Iterable[str] = (),
    exemplars: dict[str, list[str]] | None = None,
) -> Iterator[str]:
    system, _ = build_prompt(form_spec, combined_text, exemplars)
    return carenotes.refine_document(
        combined_text, draft_marker_text, instruction, backend, stream,
        history=history, phi_values=phi_values, acknowledged=acknowledged,
        system=system, refine_prompt_name="refine_form.txt",
    )


def render_preview(form_spec: FormSpec, field_values: dict[str, str]) -> str:
    """Human-readable rendering for display only — the marker text in
    ``draft_state`` (not this) is what gets refined, re-identified and
    exported."""
    lines = [f"#### {form_spec.title}"]
    for field in form_spec.fields:
        lines.append(f"\n**{field.label}**\n")
        lines.append(field_values.get(field.key) or "Not documented")
    return "\n".join(lines)
