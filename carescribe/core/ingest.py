"""
Text extraction for uploaded documents (PDF / DOCX / TXT).

Nothing here writes to disk. Uploaded bytes are held in memory only, which is
what keeps PHI off the filesystem.
"""

from __future__ import annotations

import io
from typing import Any

SUPPORTED_EXTENSIONS = ("pdf", "docx", "txt")


class IngestError(RuntimeError):
    """Raised when a document can't be read or contains no extractable text."""


def _read_bytes(file: Any) -> tuple[str, bytes]:
    """Normalise a Streamlit UploadedFile (or a path string) to (name, bytes)."""
    if isinstance(file, (str, bytes)) and not isinstance(file, bytes):
        with open(file, "rb") as handle:
            return file, handle.read()

    name = getattr(file, "name", "")
    if hasattr(file, "getvalue"):
        return name, file.getvalue()
    if hasattr(file, "read"):
        if hasattr(file, "seek"):
            file.seek(0)
        return name, file.read()

    raise IngestError("Unsupported file object — expected an upload or a file path.")


def _extract_pdf(data: bytes) -> str:
    """PDF text via pdfplumber, falling back to pypdf.

    pdfplumber gives better layout fidelity on clinical forms; pypdf is the
    safety net for files its parser chokes on.
    """
    try:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = "\n\n".join(pages).strip()
        if text:
            return text
    except Exception:
        pass  # fall through to pypdf

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(pages).strip()
    except Exception as exc:  # noqa: BLE001
        raise IngestError(f"Could not read the PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    """DOCX text: paragraphs plus table cells (clinical forms lean on tables)."""
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()
    except IngestError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise IngestError(f"Could not read the DOCX file: {exc}") from exc


def _extract_txt(data: bytes) -> str:
    """Plain text, trying the encodings clinical exports actually use."""
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    # latin-1 never fails, but keep a defensive final pass.
    return data.decode("utf-8", errors="replace").strip()


def extract_text(file: Any) -> str:
    """Extract plain text from an uploaded pdf/docx/txt file.

    Raises :class:`IngestError` for unsupported types, unreadable files, or
    documents with no extractable text (e.g. a scanned PDF with no OCR layer).
    """
    name, data = _read_bytes(file)

    if not data:
        raise IngestError("The uploaded file is empty.")

    extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if extension == "pdf":
        text = _extract_pdf(data)
    elif extension == "docx":
        text = _extract_docx(data)
    elif extension == "txt":
        text = _extract_txt(data)
    elif extension == "doc":
        raise IngestError(
            "Legacy .doc files aren't supported. Save the document as .docx and retry."
        )
    else:
        raise IngestError(
            f"Unsupported file type '.{extension}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}."
        )

    if not text.strip():
        raise IngestError(
            "No text could be extracted. If this is a scanned PDF, it has no text "
            "layer — OCR it first (CareScribe does not run OCR)."
        )

    return text
