"""
CareScribe — local, privacy-preserving de-identification and review.

Run with:  streamlit run carescribe/app.py

Privacy model
-------------
* CPU-only and fully offline. This stage makes no network calls of any kind:
  no LLM, no provider, no telemetry. The server binds to 127.0.0.1 (see
  ``.streamlit/config.toml``).
* Raw document text, the detected identifiers, and the identity mapping
  (real value -> placeholder) live ONLY in ``st.session_state`` — server-side
  RAM. None of it is written to disk, ever.
* The single write path is approval, and it writes de-identified text only, to
  ``carescribe/output/deidentified/``.

Generation is deliberately not wired up. See :mod:`carescribe.core.carenotes`.
"""

from __future__ import annotations

import html
import io
import sys
import time
from pathlib import Path

import pandas as pd
import streamlit as st

# Allow `streamlit run carescribe/app.py` to resolve the `carescribe` package
# by putting its parent directory on sys.path.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from carescribe.core import (  # noqa: E402
    applog, backends, batch, carenotes, deidentify, desktop, generation_status,
    ingest, mapping, model_setup, ollama_client, review_checklist, review_flags,
    review_spans,
)
from carescribe.components.highlight_review import highlight_review  # noqa: E402
from carescribe.ui import components as ui, theme as ui_theme  # noqa: E402

st.set_page_config(page_title="CareScribe", page_icon="🩺", layout="wide")

TABLE_COLUMNS = ["value", "type", "placeholder", "action", "confidence"]


# --------------------------------------------------------------------------
# Session state
#
# Every key here can hold PHI or PHI-derived data. "Wipe PHI" clears exactly
# this list, so a new PHI-bearing key means a new entry here too.
# --------------------------------------------------------------------------

PHI_KEYS: dict = {
    "docs": {},        # filename -> batch.Document (raw text, entities, map)
    "order": [],       # filenames, in batch order
    "selected": "",    # filename currently under review
    "load_errors": [],
    # Reviewer decisions, keyed by filename. A dismissal key carries the text of
    # the span that was dismissed — which is precisely a string someone decided
    # was not an identifier, and could be wrong. It belongs on this list.
    "entity_confirmed": {}, # filename -> set of confirmed low-confidence entity value-keys
    "flag_dismissed": {},   # filename -> list of dismissed flag keys
    "flag_redacted": {},    # filename -> count redacted from highlights
    # Generated drafts. The de-identified draft is safe; the re-identified one
    # holds real identifiers, so it lives here and is wiped with everything else.
    "drafts": {},
    # Clinical-form drafts (multi-document, keyed by selection+form). Holds the
    # merged identity map, re-identified field values, and typed header values
    # (client name, DOB, session number) — all real PHI.
    "form_drafts": {},
}

DEFAULTS = {**PHI_KEYS, "uploader_nonce": 0, "folder_path": ""}


def init_state() -> None:
    for key, value in DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value.copy() if isinstance(value, (dict, list)) else value


def wipe_phi() -> None:
    """Drop every document, identifier table, and identity map from memory."""
    for key, value in PHI_KEYS.items():
        st.session_state[key] = value.copy() if isinstance(value, (dict, list)) else value
    # Clinical-form header inputs are dynamically keyed per draft/header
    # (f"hdr_{draft_key}_{header.key}"), so they can't be listed statically
    # in PHI_KEYS — each one holds real, typed PHI (client name, DOB, ...).
    for key in list(st.session_state.keys()):
        if key.startswith(("hdr_", "attest_", "textbox_ack_")):
            del st.session_state[key]
    # Not PHI themselves, but stale UI state after a wipe.
    st.session_state.pop("form_type", None)
    st.session_state.pop("form_sources", None)
    st.session_state.pop("_batch_approve_summary", None)
    # Force the uploader to forget its files by rotating its widget key.
    st.session_state.uploader_nonce = st.session_state.get("uploader_nonce", 0) + 1


init_state()


@st.cache_resource(show_spinner=False)
def load_detection_engine() -> dict:
    """Load the NER model once per session, not once per rerun.

    Streamlit re-runs the whole script on every interaction. Without this cache
    the engine would be rebuilt on each click, which on a weak laptop is the
    difference between a responsive app and one that appears to hang every time
    the user touches anything.
    """
    import time as _time

    started = _time.monotonic()
    engine = deidentify.get_analyzer()
    status = deidentify.engine_status()
    return {
        "engine": engine,
        "model": status.get("ner_model"),
        "error": status.get("ner_error"),
        "elapsed": _time.monotonic() - started,
    }


def ensure_engine_ready() -> dict:
    """Load the model at startup, behind a visible spinner.

    Deliberately not lazy. If the first "De-identify" click is what triggers a
    multi-second load, the click looks like it did nothing.
    """
    if "engine_state" in st.session_state:
        return st.session_state["engine_state"]

    with st.spinner(
        "Loading the de-identification model — first start can take up to a "
        "minute on this computer."
    ):
        state = load_detection_engine()
    st.session_state["engine_state"] = state

    if state["error"]:
        applog.warn("detection engine unavailable at startup")
    else:
        applog.log(
            "detection engine ready model=%s elapsed=%.1fs",
            state["model"], state["elapsed"],
        )
    return state


def render_engine_failure(state: dict) -> None:
    """A missing model must stop loudly, never fall back to fetching one."""
    st.error(
        "**De-identification is not available — the language model could not "
        "be loaded.**\n\nNothing has been sent anywhere and no document has "
        "been changed."
    )
    st.code(state["error"] or "unknown error", language="text")
    st.caption(f"Full detail is in the log: `{applog.log_path()}`")


def documents() -> dict[str, batch.Document]:
    return st.session_state.docs


def current() -> batch.Document | None:
    return documents().get(st.session_state.selected)


def refresh(document: batch.Document, entities: list[dict]) -> None:
    """Re-derive the preview and map from an edited entity list."""
    result = deidentify.rebuild(document.raw_text, entities)
    document.entities = result.entities
    document.redacted_text = result.redacted_text
    document.phi_map = result.phi_map
    document.known_as = result.known_as
    # Any edit invalidates a previous approval decision.
    document.approved = False
    document.residual = []


# --------------------------------------------------------------------------
# Sidebar
# --------------------------------------------------------------------------

def _render_generation_model() -> None:
    """Name the generation model, and show its model card if one ships with it."""
    from carescribe.core import desktop

    model_path = desktop.find_local_model()
    if model_path is None:
        return
    st.sidebar.subheader("Generation model")
    st.sidebar.markdown(
        f'<div class="cs-layer" data-state="on">{ui.icon("cpu")}'
        f'<span class="cs-mono" style="font-size:.8rem">{model_path.stem}</span></div>',
        unsafe_allow_html=True,
    )

    stem = model_path.name.split(".", 1)[0]
    card = model_path.parent / f"{stem}.MODEL_CARD.md"
    if not card.is_file():
        card = model_path.parent / "MODEL_CARD.md"
    if card.is_file():
        with st.sidebar.expander("Model card", expanded=False):
            st.markdown(card.read_text(encoding="utf-8"))


def render_sidebar() -> None:
    st.sidebar.markdown(
        f'<h1 style="display:flex;align-items:center;gap:.5rem;margin:.2rem 0 .1rem">'
        f'<span style="color:var(--cs-accent);display:inline-flex">{ui.icon("shield")}'
        f'</span>CareScribe</h1>',
        unsafe_allow_html=True,
    )
    with st.sidebar:
        privacy_indicator()

    st.sidebar.subheader("Detection layers")
    status = deidentify.engine_status()

    layers = [ui.detection_layer("on", 1, "Structured regex", "always on")]
    if status["ner"]:
        layers.append(ui.detection_layer("on", 2, "Presidio + spaCy", status["ner_model"]))
    elif status["ner_error"]:
        layers.append(ui.detection_layer("warn", 2, "Presidio + spaCy", "unavailable"))
    else:
        layers.append(ui.detection_layer("wait", 2, "Presidio + spaCy", "loads on first document"))
    if status["gliner"]:
        layers.append(ui.detection_layer("on", 3, "GLiNER", "loaded"))
    else:
        layers.append(ui.detection_layer("off", 3, "GLiNER", "not installed (optional)"))
    st.sidebar.markdown("".join(layers), unsafe_allow_html=True)
    if status["ner_error"]:
        st.sidebar.caption(status["ner_error"])

    st.sidebar.caption(
        "In-prose dates: "
        + ("redacted" if status["inprose_dates"] else "kept unless identity-anchored")
    )

    _render_generation_model()

    st.sidebar.divider()

    st.sidebar.subheader("Session")
    docs = documents()
    approved = sum(1 for doc in docs.values() if doc.approved)
    identifiers = sum(len(doc.entities) for doc in docs.values())
    st.sidebar.markdown(
        ui.stat_strip([
            ("Documents in memory", len(docs)),
            ("Identifiers detected", identifiers),
            ("Approved", f"{approved} / {len(docs)}" if docs else "0"),
        ]),
        unsafe_allow_html=True,
    )
    st.sidebar.write("")

    if st.sidebar.button(
        "Clear session — wipe PHI", type="primary",
        use_container_width=True, key="wipe_phi_btn",
    ):
        wipe_phi()
        st.rerun()

    st.sidebar.caption(
        "Wipe drops every document, identifier table, and identity map from "
        "memory. None of it was ever written to disk. Approved de-identified "
        "files already on disk are left alone."
    )


# --------------------------------------------------------------------------
# 1. Load a batch
# --------------------------------------------------------------------------

def ingest_sources(sources: list) -> None:
    """Extract text from uploads/paths into session state."""
    with st.spinner(f"Reading {len(sources)} file(s)…"):
        loaded, errors = batch.load_documents(sources)

    st.session_state.docs = loaded
    st.session_state.order = list(loaded)
    st.session_state.selected = st.session_state.order[0] if st.session_state.order else ""
    st.session_state.load_errors = errors


def section_load() -> None:
    st.subheader("1. Load a batch")
    st.caption(
        "Files are read into memory only. Nothing is copied to disk at this "
        "stage, and the originals are never modified."
    )

    upload_tab, folder_tab = st.tabs(["Upload files", "Point at a local folder"])

    with upload_tab:
        nonce = st.session_state.uploader_nonce
        uploaded = st.file_uploader(
            "Choose PDF, DOCX, or TXT files",
            type=list(ingest.SUPPORTED_EXTENSIONS),
            accept_multiple_files=True,
            key=f"uploader_{nonce}",
        )
        if uploaded and st.button("Load uploaded files", type="primary"):
            ingest_sources(list(uploaded))
            st.rerun()

    with folder_tab:
        folder = st.text_input(
            "Folder path",
            key="folder_path",
            placeholder=r"C:\Users\you\Documents\batch",
            help="Read non-recursively. Files already ending in .deid.txt are skipped.",
        )
        if st.button("Load folder", disabled=not folder):
            try:
                paths = batch.list_folder(folder)
            except batch.BatchError as exc:
                st.error(str(exc))
            else:
                ingest_sources([str(p) for p in paths])
                st.rerun()

    for message in st.session_state.load_errors:
        st.warning(message)

    docs = documents()
    if docs:
        st.success(f"Loaded **{len(docs)}** document(s) into memory.")
        oversized = [
            name for name, doc in docs.items()
            if len(doc.raw_text) > deidentify.SOFT_CHAR_LIMIT
        ]
        if oversized:
            st.info(
                "Long document(s) — expect a slower review pass: "
                + ", ".join(f"`{name}`" for name in oversized)
            )


# --------------------------------------------------------------------------
# 2. Process the batch, one document at a time
# --------------------------------------------------------------------------

def section_process() -> None:
    docs = documents()
    order = st.session_state.order
    if not docs:
        return

    st.subheader("2. De-identify")

    pending = [name for name in order if not docs[name].analyzed]
    st.caption(
        "Documents are processed one at a time on the CPU. The first one also "
        "pays for loading the spaCy model."
    )

    if pending and st.button(
        f"Run de-identification on {len(pending)} document(s)", type="primary"
    ):
        progress = st.progress(0.0)
        label = st.empty()
        total = len(pending)
        for index, name in enumerate(pending, start=1):
            label.write(f"Document {index} of {total} — `{name}`")
            progress.progress((index - 1) / total)
            batch.analyze_document(docs[name])
            progress.progress(index / total)
        label.write(f"Processed {total} document(s).")
        st.rerun()

    analyzed = sum(1 for doc in docs.values() if doc.analyzed)
    st.progress(analyzed / len(order) if order else 0.0)
    st.caption(f"{analyzed} of {len(order)} documents analysed.")

    failures = [doc for doc in docs.values() if doc.error]
    for doc in failures:
        st.error(f"`{doc.name}`: {doc.error}")


# --------------------------------------------------------------------------
# 3. Review & approve, per document
# --------------------------------------------------------------------------

def entity_frame(document: batch.Document) -> pd.DataFrame:
    rows = [
        {
            "value": entity.get("value", ""),
            "type": entity.get("type", "OTHER_ID"),
            "placeholder": entity.get("placeholder", ""),
            "action": mapping.normalise_action(entity.get("action")),
            "confidence": str(entity.get("confidence") or "review"),
        }
        for entity in document.entities
    ]
    frame = pd.DataFrame(rows, columns=TABLE_COLUMNS)
    # Single-detector ("review") rows first: those are the ones most worth a
    # human's eye, in either direction. Stable, so order is otherwise unchanged.
    if not frame.empty:
        frame = frame.sort_values(
            by="confidence", key=lambda col: col.ne("review"), kind="stable"
        ).reset_index(drop=True)
    return frame


def render_entity_table(document: batch.Document) -> None:
    st.markdown("#### Detected identifiers")
    st.caption(
        "Fix a type, correct a value, delete a false positive, or switch an "
        "action to **Keep** to leave that string in the document."
    )

    redacted = [
        entity for entity in document.entities
        if mapping.normalise_action(entity.get("action")) == mapping.REDACT
    ]
    single_layer = sum(
        1 for entity in redacted
        if str(entity.get("confidence") or "review") == "review"
    )
    if single_layer:
        st.caption(
            f"⚠️ {single_layer} of {len(redacted)} redaction(s) were flagged by "
            "a single detector — the likeliest to be wrong either way. They are "
            "listed first here and highlighted in the preview above."
        )

    edited = st.data_editor(
        entity_frame(document),
        num_rows="dynamic",
        use_container_width=True,
        height=420,
        key=f"editor_{document.name}",
        column_config={
            "value": st.column_config.TextColumn(
                "Detected text", required=True, width="large",
                help="Must match the document verbatim. Line breaks inside a "
                     "name are handled for you.",
            ),
            "type": st.column_config.SelectboxColumn(
                "Type", options=list(mapping.ENTITY_TYPES), required=True, width="medium"
            ),
            "placeholder": st.column_config.TextColumn(
                "Placeholder", width="medium",
                help="Leave blank to auto-assign, e.g. [CLINICIAN_1].",
            ),
            "action": st.column_config.SelectboxColumn(
                "Action", options=[mapping.REDACT, mapping.KEEP], required=True, width="small",
                help="Keep leaves the text in place — use it for a false positive "
                     "you would rather not delete from the table.",
            ),
            "confidence": st.column_config.TextColumn(
                "Found by", disabled=True, width="small",
                help="'auto' — corroborated by more than one layer, or a "
                     "structured format. 'review' — a single detector, so worth "
                     "a second look.",
            ),
        },
    )

    if st.button("Apply table edits", use_container_width=True, key=f"apply_{document.name}"):
        # 'confidence' is display-only; drop it so rebuild re-derives tiering.
        records = [
            {key: value for key, value in row.items() if key != "confidence"}
            for row in edited.fillna("").to_dict("records")
        ]
        refresh(document, records)
        st.rerun()


def render_add_missed(document: batch.Document) -> None:
    with st.expander("Add a missed identifier", expanded=False):
        st.caption(
            "Anything the layers missed. It is variant-expanded like a detected "
            "value, so a full name also covers the title+surname and initials forms."
        )

        value_column, type_column, button_column = st.columns([3, 2, 1])
        with value_column:
            value = st.text_input(
                "Value", key=f"missed_value_{document.name}", label_visibility="collapsed",
                placeholder="Paste the exact text as it appears in the document",
            )
        with type_column:
            entity_type = st.selectbox(
                "Type", list(mapping.ENTITY_TYPES), key=f"missed_type_{document.name}",
                label_visibility="collapsed",
            )
        with button_column:
            add = st.button("Add", use_container_width=True, key=f"missed_add_{document.name}")

    if add:
        try:
            result = deidentify.add_manual_entity(
                document.raw_text, document.entities, value, entity_type
            )
        except deidentify.DeidentificationError as exc:
            st.error(str(exc))
        else:
            document.entities = result.entities
            document.redacted_text = result.redacted_text
            document.phi_map = result.phi_map
            document.approved = False
            document.residual = []
            if value.strip() and value.strip() not in document.raw_text:
                st.warning(
                    f"'{value.strip()}' was added, but that exact string does not "
                    "appear in the document — check spelling and spacing."
                )
            st.rerun()


def render_coverage(document: batch.Document) -> None:
    expanded = mapping.surface_forms(document.entities, document.known_as)

    with st.expander(
        f"Surface forms covered — {len(expanded.forms)} strings from "
        f"{len(document.entities)} rows",
        expanded=False,
    ):
        st.caption(
            "Each row also redacts these derived forms: title+surname, initials, "
            "first+surname, and organisation short forms. Every variant of one "
            "person maps to that person's single placeholder."
        )
        for placeholder, forms in expanded.by_placeholder.items():
            st.markdown(f"**`{placeholder}`** — " + ", ".join(f"`{f}`" for f in sorted(forms)))

    if expanded.ambiguous:
        st.warning(
            "These forms are claimed by more than one identifier and were "
            "assigned to the first: "
            + ", ".join(f"`{form}` → `{kept}`" for form, kept, _ in expanded.ambiguous[:10])
            + ". They are still redacted."
        )

    leftovers = mapping.residual_values(
        document.redacted_text, document.entities, document.known_as
    )
    if leftovers:
        st.warning(
            "These table values still appear in the preview — check for spelling "
            "differences against the document: "
            + ", ".join(f"`{value}`" for value in leftovers[:10])
        )


def write_approved_word(document: batch.Document) -> None:
    """Redact the original .docx using the map the reviewer just approved.

    Detection is not re-run. The map comes from the approved entity table, so
    the Word file cannot drift from the text the reviewer signed off on. The
    write itself re-scans the finished document and refuses it if anything
    identifying survives, so the Word path clears the same bar as the text one.
    """
    if not document.source_bytes:
        return
    try:
        document.approved_docx_path = str(
            batch.write_approved_docx(
                document.name,
                document.source_bytes,
                batch.approved_map(document.entities, document.known_as),
                acknowledged=document.dismissed,
            )
        )
    except batch.BatchError as exc:
        document.approved_docx_path = ""
        st.error(f"Word output blocked — {exc}")


def render_docx_download(document: batch.Document) -> None:
    """Offer the redacted .docx, but only once it has cleared the sweep."""
    if not document.source_bytes:
        return
    if not document.approved_docx_path:
        st.warning(
            "The Word version was not written — the safety sweep found "
            "something in it. The de-identified text above was still approved; "
            "resolve the finding and approve again to get the .docx."
        )
        return

    path = Path(document.approved_docx_path)
    if not path.exists():
        return
    st.download_button(
        "⬇ Download redacted Word document",
        data=path.read_bytes(),
        file_name=path.name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_docx_{document.name}",
    )
    st.caption(
        f"`{path.name}` keeps the original's tables, styles, headers and "
        "footers. Only the redacted document is written — never the mapping."
    )


_FLAG_TINTS = {
    review_flags.KIND_NAME: "#fff3bf",
    review_flags.KIND_ID: "#ffd8a8",
    review_flags.KIND_DATE: "#d0ebff",
    review_flags.KIND_INITIALS: "#e5dbff",
}


def document_flags(document: batch.Document) -> list:
    """Candidate residuals for this document, recomputed from current text."""
    return review_flags.candidate_residuals(document.redacted_text)


def flag_dismissals(document: batch.Document) -> list[str]:
    return st.session_state.flag_dismissed.setdefault(document.name, [])


def entity_confirmed(document: batch.Document) -> set[str]:
    return st.session_state.entity_confirmed.setdefault(document.name, set())


def _review_span_style(span: review_spans.ReviewSpan) -> str:
    if span.kind == review_spans.KIND_ENTITY:
        # Already redacted, already safe — a dotted underline says "worth a
        # second look", not "this is exposed text", which the solid tints
        # below correctly reserve for residual (never-redacted) candidates.
        return "border-bottom:2px dotted #868e96;padding:0 1px"
    tint = _FLAG_TINTS.get(span.flag_kind, "#f1f3f5")
    return f"background:{tint};padding:0 2px;border-radius:2px"


def _render_review_html(document: batch.Document, spans: list) -> str:
    text = document.redacted_text
    parts: list[str] = []
    cursor = 0
    for span in spans:
        parts.append(html.escape(text[cursor : span.char_start]))
        parts.append(
            f'<mark data-span-id="{html.escape(span.id)}" '
            f'style="{_review_span_style(span)}" '
            f'title="{html.escape(span.why)}">'
            f'{html.escape(text[span.char_start : span.char_end])}</mark>'
        )
        cursor = span.char_end
    parts.append(html.escape(text[cursor:]))
    return "".join(parts)


def render_review(document: batch.Document) -> list:
    """The primary review surface: the redacted text, second-look items marked.

    Returns the span list so the caller can drive the bulk actions and the
    Approve gate without recomputing it.
    """
    spans = review_spans.review_spans(
        document.redacted_text, document.entities,
        entity_confirmed(document), flag_dismissals(document),
    )

    st.markdown("#### Redacted preview")
    st.caption(
        "This is exactly what approval writes to disk. Highlighted words are "
        "optional second-look items — click one to act on it on its own, or use "
        "the bulk buttons below."
    )

    clicked = highlight_review(
        _render_review_html(document, spans), key=f"hl_{document.name}"
    )
    _render_span_action(document, spans, clicked)

    # The interactive view above is the working copy; this is the plain,
    # authoritative, copy-pasteable text — kept unstyled and disabled on
    # purpose so it's never mistaken for editable.
    with st.expander("Plain text — exactly what gets written"):
        st.text_area(
            "Redacted", document.redacted_text, height=220,
            label_visibility="collapsed", disabled=True,
            key=f"preview_exact_{document.name}",
        )

    with st.expander("Full detected-identifier table"):
        render_entity_table(document)

    return spans


def render_bulk_actions(document: batch.Document, spans: list) -> None:
    """Clear each class of second-look item in one click.

    Low-confidence entity spans are already redacted, so "Confirm all" only
    records that the reviewer accepts the placeholders. "Redact all flagged"
    over-redacts every permissive residual candidate — the safe direction.
    Either is optional: neither gates Approve.
    """
    entity_keys: list[str] = []
    residual_spans: list = []
    seen_e: set[str] = set()
    seen_r: set[str] = set()
    for span in spans:
        key = span.id.split(":", 1)[1]
        if span.kind == review_spans.KIND_ENTITY:
            if key not in seen_e:
                seen_e.add(key)
                entity_keys.append(key)
        elif key not in seen_r:
            seen_r.add(key)
            residual_spans.append(span)

    n_e, n_r = len(entity_keys), len(residual_spans)
    if not n_e and not n_r:
        st.caption("Nothing needs a second look in this document.")
        return

    bits = []
    if n_e:
        bits.append(f"**{n_e}** low-confidence redaction(s) already in place")
    if n_r:
        bits.append(f"**{n_r}** item(s) flagged by the safety sweep")
    st.markdown(" · ".join(bits))

    left, right = st.columns(2)
    with left:
        if n_e and st.button(
            f"Confirm all {n_e} redaction(s)",
            key=f"confirm_all_{document.name}", use_container_width=True,
        ):
            entity_confirmed(document).update(entity_keys)
            st.rerun()
    with right:
        if n_r and st.button(
            f"Redact all {n_r} flagged item(s)",
            key=f"redact_all_{document.name}", use_container_width=True,
        ):
            added = 0
            for span in residual_spans:
                try:
                    result = deidentify.add_manual_entity(
                        document.raw_text, document.entities, span.text
                    )
                except deidentify.DeidentificationError:
                    continue
                document.entities = result.entities
                document.redacted_text = result.redacted_text
                document.phi_map = result.phi_map
                added += 1
            if added:
                document.approved = False
                document.residual = []
                st.session_state.flag_redacted[document.name] = (
                    st.session_state.flag_redacted.get(document.name, 0) + added
                )
            st.rerun()


def _render_span_action(document: batch.Document, spans: list, clicked_id: str | None) -> None:
    if not clicked_id:
        return
    match = next((s for s in spans if s.id == clicked_id), None)
    if match is None:
        return  # already resolved by an earlier click this session

    st.markdown(f"**`{match.text}`** — {match.why}")

    if match.kind == review_spans.KIND_ENTITY:
        value_key = clicked_id.split(":", 1)[1]
        confirm_col, undo_col, retype_col = st.columns([1, 1, 2])
        with confirm_col:
            if st.button("✅ Confirm", key=f"conf_{document.name}_{value_key}"):
                entity_confirmed(document).add(value_key)
                st.rerun()
        with undo_col:
            if st.button("↩ Undo (restore text)", key=f"undo_{document.name}_{value_key}"):
                for entity in document.entities:
                    if entity["value"].strip().casefold() == value_key:
                        entity["action"] = mapping.KEEP
                refresh(document, document.entities)
                st.rerun()
        with retype_col:
            new_type = st.selectbox(
                "Change type", list(mapping.ENTITY_TYPES),
                key=f"retype_{document.name}_{value_key}", label_visibility="collapsed",
            )
            if st.button("Apply type", key=f"retype_go_{document.name}_{value_key}"):
                for entity in document.entities:
                    if entity["value"].strip().casefold() == value_key:
                        entity["type"] = new_type
                refresh(document, document.entities)
                st.rerun()
    else:
        flag_key = clicked_id.split(":", 1)[1]
        redact_col, dismiss_col = st.columns(2)
        with redact_col:
            if st.button("Redact this", key=f"resred_{document.name}_{flag_key}"):
                try:
                    result = deidentify.add_manual_entity(
                        document.raw_text, document.entities, match.text
                    )
                except deidentify.DeidentificationError as exc:
                    st.error(str(exc))
                else:
                    document.entities = result.entities
                    document.redacted_text = result.redacted_text
                    document.phi_map = result.phi_map
                    document.approved = False
                    st.session_state.flag_redacted[document.name] = (
                        st.session_state.flag_redacted.get(document.name, 0) + 1
                    )
                    st.rerun()
        with dismiss_col:
            if st.button("Not an identifier", key=f"resdis_{document.name}_{flag_key}"):
                flag_dismissals(document).append(flag_key)
                st.rerun()


def render_approval(document: batch.Document, spans: list) -> None:
    st.divider()
    st.warning(
        "**Human review required.** Automated de-identification is not a "
        "guarantee. Read the redacted preview above before approving."
    )

    n_advisory = len({s.id for s in spans})

    if document.residual:
        st.error(
            "**Approval blocked — the safety sweep found what still look like "
            "identifiers:**\n\n"
            + "\n".join(f"- `{value}`" for value in document.residual)
        )
        st.caption(
            "Handle each one: add it as a missed identifier above so it gets "
            "redacted, or — if it is genuinely not identifying (a town used as a "
            "place of care, a ward name) — dismiss it. Dismissals apply to this "
            "document only and are not saved anywhere."
        )
        for value in document.residual:
            if st.button(
                f"Dismiss `{value}` — not an identifier",
                key=f"dismiss_{document.name}_{value}",
            ):
                document.dismissed.append(value)
                document.residual = batch.sweep(document.redacted_text, document.dismissed)
                st.rerun()

    if document.dismissed:
        st.caption("Dismissed this document: " + ", ".join(f"`{v}`" for v in document.dismissed))

    if document.approved:
        st.success(f"Approved. Written to `{document.approved_path}`")
        render_docx_download(document)

    # A .docx whose text boxes this pass cannot reach must not slip through on
    # the reviewer's assumption that the preview covered everything. This is
    # the one acknowledgement this redesign keeps as an explicit checkbox
    # rather than folding into the click-to-redact model — text-box content
    # isn't a text span at all, so nothing in review_spans can ever point at
    # it. It gates Approve directly, below.
    if document.has_text_boxes:
        st.error(
            "**This document contains text boxes, shapes, or embedded objects.** "
            "Their contents are **not** redacted automatically and do not appear "
            "in the preview above. Open the original and check them by hand."
        )
        document.text_boxes_acknowledged = st.checkbox(
            "I have opened the original and checked any text boxes / shapes by hand.",
            value=document.text_boxes_acknowledged,
            key=f"textbox_ack_{document.name}",
        )

    # An explicit, recorded attestation that a human read the redacted text.
    # Automated de-identification has no way to know whether the preview was
    # actually looked at; this is the one place the reviewer says so, and it is
    # written (as a bool) into the audit sidecar. It gates Approve, below.
    if not document.approved:
        document.attested = st.checkbox(
            "I have read the redacted preview above and confirm it is safe to release.",
            value=document.attested,
            key=f"attest_{document.name}",
        )

    # Only the authoritative sweep blocks. Advisory spans (already-redacted
    # low-confidence entities, permissive residual flags) never gate the write —
    # batch.write_approved re-runs the sweep and refuses regardless.
    reason = review_checklist.blocking_reason(document.residual)
    if not reason and document.has_text_boxes and not document.text_boxes_acknowledged:
        reason = "Confirm the text-box check above first."
    if not reason and not document.attested:
        reason = "Tick the read-and-confirmed box above."

    if not reason and n_advisory:
        st.caption(
            f"{n_advisory} optional second-look item(s) not actioned — the "
            "safety sweep runs on approve and will refuse any real identifier."
        )

    if st.button(
        "✅ Run safety sweep and approve",
        type="primary",
        disabled=bool(reason) or not document.redacted_text,
        key=f"approve_{document.name}",
    ):
        document.residual = batch.sweep(document.redacted_text, document.dismissed)
        if document.residual:
            document.approved = False
        else:
            try:
                path = batch.write_approved(
                    document.name, document.redacted_text, acknowledged=document.dismissed
                )
            except batch.BatchError as exc:
                st.error(str(exc))
            else:
                document.approved = True
                document.approved_path = str(path)
                write_approved_word(document)
                batch.write_review_record(
                    document.name,
                    entities=document.entities,
                    flags_shown=len(document_flags(document)),
                    flags_redacted=st.session_state.flag_redacted.get(document.name, 0),
                    flags_dismissed=len(flag_dismissals(document)),
                    attested=document.attested,
                )
        st.rerun()

    if reason:
        st.caption(f"Approve is disabled — {reason}")

    st.caption(
        f"Approved files are written to `{batch.OUTPUT_DIR}` as de-identified "
        "text only. The identity mapping stays in memory and is never saved."
    )


def section_review() -> None:
    docs = documents()
    order = st.session_state.order
    if not docs:
        return

    ready = [name for name in order if docs[name].analyzed and not docs[name].error]
    if not ready:
        return

    st.subheader("3. Review & approve")

    if st.session_state.selected not in ready:
        st.session_state.selected = ready[0]

    def label(name: str) -> str:
        doc = docs[name]
        position = order.index(name) + 1
        mark = "✅" if doc.approved else ("⛔" if doc.residual else "🔵")
        return f"{mark} {position}/{len(order)} — {name}"

    # The index is computed rather than keyed, so the selector survives a
    # document being removed or the batch being reloaded.
    st.session_state.selected = st.radio(
        "Document under review",
        ready,
        index=ready.index(st.session_state.selected),
        format_func=label,
    )

    document = docs[st.session_state.selected]
    st.markdown(f"### `{document.name}`")
    st.caption(
        f"Document {order.index(document.name) + 1} of {len(order)} — "
        f"{len(document.raw_text):,} characters, {len(document.entities)} identifiers detected."
    )

    spans = render_review(document)
    render_bulk_actions(document, spans)
    render_add_missed(document)
    render_coverage(document)

    with st.expander("View original text (contains PHI)", expanded=False):
        st.text_area(
            "Original", document.raw_text, height=300,
            label_visibility="collapsed", disabled=True,
            key=f"raw_{document.name}",
        )

    render_approval(document, spans)
    render_batch_approve(docs, ready)


def render_batch_approve(docs: dict, ready: list) -> None:
    """One click to approve every document that comes back clean.

    Each is re-run through the blocking safety sweep; only the clean ones are
    written. Anything with sweep findings or an unchecked text-box warning is
    left for the reviewer to open individually.
    """
    pending = [docs[name] for name in ready if not docs[name].approved]

    # Show the outcome of the last batch run even if that run cleared enough
    # documents that the button below no longer renders.
    summary = st.session_state.pop("_batch_approve_summary", None)
    if summary and (summary["approved"] or summary["blocked"]):
        st.divider()
        st.markdown("#### Batch approval")
        if summary["approved"]:
            st.success(
                f"Approved {len(summary['approved'])}: "
                + ", ".join(f"`{n}`" for n in summary["approved"])
            )
        if summary["blocked"]:
            st.warning(
                "Open these individually — the safety sweep still has findings, "
                "the read-and-confirmed box isn't ticked, or a text-box check "
                "is outstanding:\n"
                + "\n".join(f"- `{n}`" for n in summary["blocked"])
            )

    if len(pending) < 2:
        return

    if not (summary and (summary["approved"] or summary["blocked"])):
        st.divider()
    st.markdown("#### Approve the whole batch")

    st.caption(
        f"{len(pending)} document(s) not yet approved. Each is re-checked by the "
        "safety sweep; only the ones that come back clean, with the "
        "read-and-confirmed box ticked, are written to disk."
    )
    if st.button(
        f"✅ Approve all {len(pending)} that pass the safety sweep",
        type="primary", key="approve_batch",
    ):
        approved: list[str] = []
        blocked: list[str] = []
        for doc in pending:
            if not doc.attested:
                blocked.append(doc.name)
                continue
            if doc.has_text_boxes and not doc.text_boxes_acknowledged:
                blocked.append(doc.name)
                continue
            doc.residual = batch.sweep(doc.redacted_text, doc.dismissed)
            if doc.residual:
                doc.approved = False
                blocked.append(doc.name)
                continue
            try:
                path = batch.write_approved(
                    doc.name, doc.redacted_text, acknowledged=doc.dismissed
                )
            except batch.BatchError:
                blocked.append(doc.name)
                continue
            doc.approved = True
            doc.approved_path = str(path)
            write_approved_word(doc)
            batch.write_review_record(
                doc.name,
                entities=doc.entities,
                flags_shown=len(document_flags(doc)),
                flags_redacted=st.session_state.flag_redacted.get(doc.name, 0),
                flags_dismissed=len(flag_dismissals(doc)),
                attested=doc.attested,
            )
            approved.append(doc.name)
        st.session_state["_batch_approve_summary"] = {
            "approved": approved, "blocked": blocked,
        }
        st.rerun()


# --------------------------------------------------------------------------
# 4. Batch status
# --------------------------------------------------------------------------

def section_batch_status() -> None:
    docs = documents()
    order = st.session_state.order
    if not docs:
        return

    st.subheader("4. Batch status")

    def _kind(doc: batch.Document) -> str:
        if doc.error:
            return "failed"
        if doc.approved:
            return "approved"
        if doc.residual:
            return "blocked"
        if doc.analyzed:
            return "review"
        return "pending"

    body = []
    for position, name in enumerate(order, start=1):
        doc = docs[name]
        out = html.escape(doc.approved_path) if doc.approved_path else "&mdash;"
        body.append(
            "<tr>"
            f'<td>{position}</td>'
            f'<td><span class="cs-mono">{html.escape(name)}</span></td>'
            f'<td>{ui.status_chip(_kind(doc))}</td>'
            f'<td>{len(doc.entities)}</td>'
            f'<td><span class="cs-mono">{out}</span></td>'
            "</tr>"
        )
    st.markdown(
        '<table class="cs-table"><thead><tr>'
        "<th>#</th><th>Document</th><th>Status</th>"
        "<th>Identifiers</th><th>Approved output</th>"
        f'</tr></thead><tbody>{"".join(body)}</tbody></table>',
        unsafe_allow_html=True,
    )

    approved = sum(1 for doc in docs.values() if doc.approved)
    st.caption(f"{approved} of {len(order)} approved.")


# --------------------------------------------------------------------------
# 5. Handoff — deliberately not wired up
# --------------------------------------------------------------------------

def _draft_state(name: str) -> dict:
    return st.session_state.drafts.setdefault(
        name, {"deidentified": "", "reidentified": "", "history": [], "unresolved": []}
    )


def _form_draft_key(document_names: list[str], form_id: str) -> str:
    return "|".join(sorted(document_names)) + "::" + form_id


def _form_draft_state(key: str) -> dict:
    return st.session_state.setdefault("form_drafts", {}).setdefault(
        key,
        {
            "deidentified": "",       # marker text — refine/reidentify/export source of truth
            "unresolved": [],
            "history": [],
            "field_values": {},       # parsed {field_key: text}, deidentified
        },
    )


def _invalidate_form_export(draft: dict) -> None:
    """Drop any previously re-identified/exportable content — called
    whenever the underlying draft text changes (fresh generation or a
    refinement), so a stale export can never be served after the content
    it was built from is gone."""
    draft.pop("resolved_field_values", None)
    draft["unresolved"] = []


def _header_values_complete(form_spec, header_values: dict) -> bool:
    required = [h for h in form_spec.header_fields if h.key != "reason_for_referral"]
    return all((header_values.get(h.key) or "").strip() for h in required)


def privacy_indicator() -> None:
    """A persistent, honest statement of where data goes.

    It must change when cloud generation is configured. An indicator that says
    "fully offline" while text is leaving the machine is worse than no
    indicator, because it is the thing a clinician would point at.
    """
    if backends.cloud_enabled():
        st.warning(
            f"**Cloud generation is enabled ({backends.cloud_provider()}).** "
            "De-identification and review happen entirely on this computer. "
            "For generation, the **approved de-identified text** — placeholders "
            "only, never real identifiers — is sent to "
            f"{backends.cloud_provider()}. Generation still requires your "
            "approval and a clean safety sweep first.",
            icon=":material/cloud:",
        )
    elif st.session_state.get("downloading_model"):
        st.info(
            "**Downloading the AI model onto this computer.** Weights are "
            "coming *in*; no patient data is going out. Nothing about any "
            "document is part of this request.",
            icon=":material/download:",
        )
    else:
        st.success(
            "**Running fully offline — no data leaves this computer.** "
            "Documents are read into memory, de-identified, reviewed and "
            "drafted here. Nothing is uploaded.",
            icon=":material/lock:",
        )


def render_generation_status() -> dict:
    """Which backend will be used, and the fix if none is available."""
    state = backends.describe_backends()

    verdict = desktop.ram_verdict()
    if not verdict["ok"]:
        st.warning(verdict["message"])

    if state["ollama"]["available"]:
        st.success(
            f"Generating with **Ollama · {state['ollama']['default_model']}** "
            "on this computer."
        )
    elif state["local"]["available"]:
        name = Path(state["local"]["model_path"]).name
        st.success(f"Generating with the **built-in model** (`{name}`) on this computer.")
        st.caption(
            "The built-in model is small enough to run on an ordinary laptop. "
            "For longer or more complex documents, installing Ollama and "
            "pulling an 8B model gives noticeably better drafts — CareScribe "
            "will use it automatically."
        )
    elif state["cloud"]["available"]:
        st.warning(
            f"No local model available. Generating via **{state['cloud']['provider']}** "
            "with de-identified text only."
        )
    else:
        st.error("No generation backend is available.")
        st.code(
            "De-identification and review still work.\n\n"
            "For generation, either reinstall CareScribe so the built-in model\n"
            "is present, or install Ollama and run:\n"
            "    ollama pull qwen2.5:3b",
            language="text",
        )
    return state


SETUP_CARD_HEADING = "Set up generation"


def render_setup_card() -> None:
    """Shown instead of an empty panel when no model is available yet.

    An empty generation panel reads as a broken app. This is the alternative:
    say what is missing, in plain words, and offer the one click that fixes it.
    """
    status = generation_status.generation_status()

    st.markdown(f"#### {SETUP_CARD_HEADING}")
    st.info(
        f"{generation_status.missing_reason(status)}\n\n"
        "**De-identification and review work now — this is only needed for "
        "drafting notes and letters.** Setting it up is a one-time step."
    )

    option_a, option_b = st.columns(2)

    with option_a:
        st.markdown("**Option A — Download the built-in model**")
        st.caption(
            f"Recommended. About "
            f"{desktop.MODEL_APPROX_BYTES / 1e9:.1f} GB, downloaded once.\n\n"
            "This downloads the AI model **onto this computer** — a one-time "
            "setup. No patient data is involved and nothing is sent anywhere."
        )
        if st.button("⬇ Download the model", type="primary", key="setup_download"):
            run_model_download()

    with option_b:
        st.markdown("**Option B — Use Ollama (better quality)**")
        if status.ollama_running:
            st.caption(
                "Ollama is running on this computer but has no model yet. "
                "Downloading one here gives noticeably better drafts."
            )
            if st.button("⬇ Download llama3.1:8b via Ollama", key="setup_pull"):
                run_ollama_pull()
        else:
            st.caption(
                "Free, installed separately, and produces better drafts than "
                "the built-in model."
            )
            st.markdown(f"[Install Ollama]({model_setup.OLLAMA_INSTALL_URL})")
            for step in model_setup.OLLAMA_STEPS:
                st.markdown(f"- {step}")
            if st.button("↻ Refresh", key="setup_refresh"):
                st.rerun()

    if status.cloud:
        st.caption(
            f"Cloud generation is configured ({status.cloud_provider}). It is "
            "off by default and only ever receives approved de-identified text."
        )


def run_model_download() -> None:
    """Option A. The only outbound request the app makes, on an explicit click."""
    st.session_state["downloading_model"] = True
    bar = st.progress(0.0, text="Starting download…")
    try:
        model_setup.download_model(
            on_progress=lambda p: bar.progress(
                p.fraction, text=f"Downloading the model — {p.message}"
            )
        )
    except model_setup.ModelSetupError as exc:
        st.session_state["downloading_model"] = False
        bar.empty()
        st.error(str(exc))
        if st.button("Discard the partial download and start over", key="setup_clear"):
            model_setup.clear_partial_download()
            st.rerun()
        return
    st.session_state["downloading_model"] = False
    bar.empty()
    st.success("Model downloaded. Generation is ready.")
    st.rerun()


def run_ollama_pull(model: str = "llama3.1:8b") -> None:
    """Option B. Ollama does the fetching; the request goes to loopback."""
    bar = st.progress(0.0, text=f"Asking Ollama for {model}…")
    try:
        for progress in model_setup.pull_ollama_model(model):
            bar.progress(progress.fraction, text=f"{model} — {progress.message}")
    except model_setup.ModelSetupError as exc:
        bar.empty()
        st.error(str(exc))
        return
    bar.empty()
    st.success(f"{model} is installed. Generation is ready.")
    st.rerun()


def render_test_generation() -> None:
    """A concrete "it works", rather than asking the clinician to trust a flag."""
    if st.button("Test generation", key="gen_selftest"):
        with st.spinner("Running a short test on this computer…"):
            try:
                _, backend, label = backends.select_backend()
                sample = "".join(
                    carenotes.generate_document(
                        "Patient: [PATIENT]\nSeen in clinic. Sertraline 50mg daily.\n",
                        "SOAP care note", backend, stream=False,
                    )
                )
            except (carenotes.CareNoteError, backends.BackendError) as exc:
                st.error(f"Generation is not working yet:\n\n{exc}")
                return
        if sample.strip():
            st.success(f"Generation is ready ✓  ({label})")
            with st.expander("What it produced"):
                st.code(sample[:600], language="markdown")
        else:
            st.error("The model returned nothing. Try the other setup option.")


def render_generation_panel(document: batch.Document) -> None:
    """Generate, refine, re-identify and export — for one approved document.

    Two dictionaries are in play and they are NOT interchangeable:

    * ``draft`` — this document's generated output. ``draft["deidentified"]`` is
      the drafted note *with placeholders still in it*.
    * ``backends_available`` — which generation backends are usable right now.

    They were both called ``state`` until one silently overwrote the other.
    """
    # Generation must never run on text a human has not approved. Reaching here
    # without approved text is a bug elsewhere, but a clinician must get a
    # sentence rather than a traceback.
    if not document.approved or not (document.redacted_text or "").strip():
        st.info(
            "This document hasn't been approved for generation yet — approve "
            "it in step 3 first."
        )
        return

    draft = _draft_state(document.name)

    if not generation_status.generation_status().ready:
        render_setup_card()
        return

    backends_available = render_generation_status()
    render_test_generation()

    template = st.selectbox(
        "Template", carenotes.template_names(), key=f"tpl_{document.name}"
    )

    custom = ""
    if template == carenotes.CUSTOM_TEMPLATE:
        custom = st.text_area(
            "Your template / instructions",
            key=f"custom_{document.name}",
            height=140,
            placeholder="Paste your house format here…",
        )

    ready = any(
        backends_available[kind]["available"] for kind in ("ollama", "local", "cloud")
    )
    if st.button(
        "✨ Generate draft",
        type="primary",
        disabled=not ready,
        key=f"gen_{document.name}",
    ):
        _run_generation(document, template, custom, draft)

    if draft.get("deidentified"):
        render_draft(document, draft)


def _stream_into(placeholder, chunks, started: float) -> str:
    """Render a stream token by token so a slow local model looks alive."""
    collected: list[str] = []
    for chunk in chunks:
        collected.append(chunk)
        placeholder.markdown(
            "".join(collected) + " ▌"
            + f"\n\n*{time.monotonic() - started:.0f}s elapsed*"
        )
    return "".join(collected)


def _run_generation(document, template, custom, draft_state) -> None:
    """First-pass generation. The model receives de-identified text only."""
    placeholder = st.empty()
    started = time.monotonic()
    try:
        with st.spinner("Generating on this computer — this can take a minute. Nothing leaves your device."):
            _, backend, _label = backends.select_backend()
            chunks = carenotes.generate_document(
                document.redacted_text,
                template,
                backend,
                stream=True,
                custom_instruction=custom,
                # Passed only so generation can assert these are absent.
                phi_values=list(document.phi_map.values()),
                # Sweep findings the reviewer cleared at approval — so the
                # pre-generation re-scan doesn't trip on them.
                acknowledged=document.dismissed,
            )
            draft = _stream_into(placeholder, chunks, started)
    except (carenotes.CareNoteError, backends.BackendError) as exc:
        placeholder.empty()
        st.error(str(exc))
        return

    placeholder.empty()
    draft_state["deidentified"] = carenotes.with_banner(draft)
    draft_state["reidentified"] = ""
    draft_state["unresolved"] = []
    draft_state["history"] = []
    st.rerun()


def render_draft(document: batch.Document, draft_state: dict) -> None:
    """The de-identified draft, refinement, re-identification, and exports."""
    st.markdown("#### Draft (de-identified)")
    st.caption(
        "Still contains placeholders — safe to display, share, and save. This "
        "is what the model produced."
    )
    st.markdown(draft_state["deidentified"])

    stem = batch.safe_stem(document.name)
    col_md, col_txt, col_docx = st.columns(3)
    with col_md:
        st.download_button(
            "⬇ .md", draft_state["deidentified"], file_name=f"{stem}.draft.md",
            mime="text/markdown", key=f"dl_md_{document.name}",
        )
    with col_txt:
        st.download_button(
            "⬇ .txt", draft_state["deidentified"], file_name=f"{stem}.draft.txt",
            mime="text/plain", key=f"dl_txt_{document.name}",
        )
    with col_docx:
        st.download_button(
            "⬇ .docx", _as_docx(draft_state["deidentified"]),
            file_name=f"{stem}.draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_docx_{document.name}",
        )

    render_refinement(document, draft_state)
    render_reidentification(document, draft_state)


def render_refinement(document: batch.Document, draft_state: dict) -> None:
    """Follow-up instructions, on de-identified text only."""
    with st.expander("Refine this draft", expanded=False):
        st.caption(
            "Refinement runs on the same de-identified source and the current "
            "draft. No real identifier enters this loop."
        )
        for instruction, _ in draft_state["history"]:
            st.markdown(f"- _{instruction}_")
        instruction = st.text_input(
            "What would you like changed?",
            key=f"refine_{document.name}",
            placeholder="e.g. make the plan more concise; add a risk paragraph",
        )
        status = ollama_client.status()
        if st.button(
            "Apply", key=f"refine_go_{document.name}",
            disabled=not instruction or not status["models"],
        ):
            placeholder = st.empty()
            started = time.monotonic()
            try:
                with st.spinner("Revising…"):
                    chunks = carenotes.refine_document(
                        document.redacted_text,
                        draft_state["deidentified"],
                        instruction,
                        backends.select_backend()[1],
                        stream=True,
                        history=draft_state["history"],
                        phi_values=list(document.phi_map.values()),
                        acknowledged=document.dismissed,
                    )
                    revised = _stream_into(placeholder, chunks, started)
            except (carenotes.CareNoteError, backends.BackendError) as exc:
                placeholder.empty()
                st.error(str(exc))
                return
            placeholder.empty()
            draft_state["history"].append((instruction, ""))
            draft_state["deidentified"] = carenotes.with_banner(revised)
            draft_state["reidentified"] = ""
            draft_state["unresolved"] = []
            st.rerun()


def render_reidentification(document: batch.Document, draft_state: dict) -> None:
    """Opt-in, local-only substitution of placeholders back to real values."""
    st.markdown("#### Re-identify for the patient record (local only)")
    st.warning(
        "**This produces a document containing real patient identifiers.** It "
        "is for saving into your own local record. It happens entirely in "
        "Python on this machine — the model never sees it, and the identity "
        "mapping is still never written to disk."
    )

    issues = mapping.check_placeholder_integrity(
        draft_state["deidentified"], list(document.phi_map)
    )
    mangled = [i for i in issues if i.kind == mapping.ISSUE_MANGLED]
    if mangled:
        st.caption(
            "Placeholders the model corrupted, repaired automatically: "
            + ", ".join(f"`{i.token}` → `{i.suggestion}`" for i in mangled)
        )

    if st.button(
        "🔓 Re-identify locally", key=f"reid_{document.name}",
        disabled=not document.phi_map,
    ):
        text, unresolved = carenotes.finalise(draft_state["deidentified"], document.phi_map)
        draft_state["reidentified"] = text
        draft_state["unresolved"] = unresolved
        st.rerun()

    if draft_state["unresolved"]:
        st.error(
            "**Blocked — these placeholders could not be resolved:** "
            + ", ".join(f"`{token}`" for token in draft_state["unresolved"])
            + "\n\nA report filed with a placeholder still in it is worse than "
            "no report. Refine the draft or correct the identifier table, then "
            "re-identify again."
        )
        return

    if draft_state["reidentified"]:
        st.success("Re-identified. Every placeholder resolved.")
        st.markdown(draft_state["reidentified"])
        stem = batch.safe_stem(document.name)
        col_md, col_docx = st.columns(2)
        with col_md:
            st.download_button(
                "⬇ .md (contains PHI)", draft_state["reidentified"],
                file_name=f"{stem}.record.md", mime="text/markdown",
                key=f"dlr_md_{document.name}",
            )
        with col_docx:
            st.download_button(
                "⬇ .docx (contains PHI)", _as_docx(draft_state["reidentified"]),
                file_name=f"{stem}.record.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dlr_docx_{document.name}",
            )


def _render_template_uploader() -> None:
    """Let a clinic add its own table-based .docx form to the selector.

    Parsing and storage are local (`template_ingest`); nothing is sent
    anywhere. A blank template carries no PHI, so it is safe to persist.
    """
    from carescribe.core import clinical_forms, template_ingest

    with st.expander("➕ Add your clinic's own template (.docx)"):
        st.caption(
            "A table-based Word form — label cells with blank cells beside them. "
            "CareScribe reads its structure on this computer; nothing leaves the device."
        )
        upload = st.file_uploader("Template file", type=["docx"], key="clinical_form_template_upload")
        if upload is None:
            return
        try:
            preview = template_ingest.parse_template_bytes(upload.getvalue(), form_id="__preview__")
        except clinical_forms.ClinicalFormError as exc:
            st.error(str(exc))
            return
        sections = len({f.key.split(".")[0] for f in preview.fields})
        st.success(
            f"Detected **{len(preview.fields)}** fields across **{sections}** section(s), "
            f"plus {len(preview.header_fields)} header field(s)."
        )
        st.write([f.label for f in preview.fields])
        if st.button("Save this template", key="clinical_form_template_save"):
            template_ingest.save_template(upload.getvalue(), upload.name)
            st.success(f"Saved “{preview.title}”. Choose it from the Form list below.")
            st.rerun()


def _render_reference_uploader() -> None:
    """Add clinic reference files (formulary, pathways, protocols) to a local
    library. They are shown to the clinician during review — never sent to the
    model — so a dose or a referral criterion is never paraphrased."""
    from carescribe.core import reference_library

    loaded = reference_library.sources()
    label = "📚 Reference material" + (f" ({len(loaded)} file(s))" if loaded else "")
    with st.expander(label):
        st.caption(
            "Formularies, care pathways, local protocols (.txt or .md, no patient "
            "data). Retrieved verbatim into the review view as an aid — not fed to "
            "the model."
        )
        for name, chunk_count in loaded:
            cols = st.columns([4, 1])
            cols[0].markdown(f"- **{name}** · {chunk_count} passage(s)")
            if cols[1].button("Remove", key=f"ref_rm_{name}"):
                reference_library.remove_file(name)
                st.rerun()
        upload = st.file_uploader(
            "Add a reference file", type=["txt", "md"], key="reference_upload"
        )
        if upload is not None and st.button("Add to library", key="reference_add"):
            try:
                reference_library.add_file(upload.name, upload.getvalue())
                st.success("Added.")
            except reference_library.ReferenceError as exc:
                st.error(str(exc))
            st.rerun()


def _render_reference_panel(spec, query: str) -> None:
    """Verbatim reference passages, retrieved per field at the granularity the
    planner chose for it. Read-only — never sent to the model."""
    from carescribe.core import reference_library, retrieval_planner

    if reference_library.is_empty() or not (query or "").strip():
        return

    plans = retrieval_planner.plan(spec, query)
    groups: list[tuple[str, list]] = []
    seen: set[tuple[str, str]] = set()
    for field_plan in plans.values():
        if not field_plan.want_reference:
            continue
        hits = reference_library.search(
            field_plan.query, k=3, granularity=field_plan.granularity
        )
        fresh = [h for h in hits if (h.source, h.text) not in seen]
        for h in fresh:
            seen.add((h.source, h.text))
        if fresh:
            groups.append((field_plan.field_label, fresh))
    if not groups:
        return

    total = sum(len(hits) for _label, hits in groups)
    with st.expander(f"📚 Relevant reference material ({total})"):
        st.caption(
            "Verbatim passages matched per field. Not sent to the model — for "
            "your judgement while reviewing."
        )
        for label, hits in groups:
            st.markdown(f"##### {label}")
            for hit in hits:
                source = f"**{hit.source}**" + (f" — {hit.heading}" if hit.heading else "")
                st.markdown(source)
                st.markdown(f"> {hit.text}")


def render_clinical_form_panel(docs: dict) -> None:
    from carescribe.core import clinical_forms

    approved = [doc for doc in docs.values() if doc.approved]
    if not approved:
        st.info("Approve at least one document in step 3 to generate a clinical form.")
        return

    _render_template_uploader()
    _render_reference_uploader()

    form_options = clinical_forms.available_forms()
    form_id = st.selectbox(
        "Form", [fid for fid, _ in form_options],
        format_func=lambda fid: dict(form_options)[fid], key="form_type",
    )
    spec = clinical_forms.get_form_spec(form_id)

    selected_names = st.multiselect(
        "Source document(s)", [doc.name for doc in approved],
        default=[approved[0].name], key="form_sources",
    )
    if not selected_names:
        st.info("Select at least one approved document.")
        return

    draft_key = _form_draft_key(selected_names, form_id)
    draft = _form_draft_state(draft_key)

    st.markdown("##### Form header")
    header_values = draft.setdefault("header_values", {})
    for header in spec.header_fields:
        widget = st.text_area if header.key == "reason_for_referral" else st.text_input
        header_values[header.key] = widget(
            header.label, value=header_values.get(header.key, ""),
            key=f"hdr_{draft_key}_{header.key}",
        )

    if not generation_status.generation_status().ready:
        render_setup_card()
        return

    backends_available = render_generation_status()
    ready = (
        any(backends_available[kind]["available"] for kind in ("ollama", "local", "cloud"))
        and _header_values_complete(spec, header_values)
    )
    if not _header_values_complete(spec, header_values):
        st.caption("Fill in every header field except Reason for referral to enable generation.")

    if st.button("✨ Generate form", type="primary", disabled=not ready, key=f"gen_form_{draft_key}"):
        _run_form_generation(docs, selected_names, spec, draft)

    if draft.get("deidentified"):
        render_form_draft(docs, selected_names, spec, draft)


def _run_form_generation(docs: dict, selected_names: list[str], spec, draft: dict) -> None:
    from carescribe.core import clinical_forms, exemplars, retrieval_planner

    sources = [(name, docs[name].redacted_text, docs[name].phi_map) for name in selected_names]
    phi_values = [v for name in selected_names for v in docs[name].phi_map.values()]
    acknowledged = [v for name in selected_names for v in docs[name].dismissed]
    combined_text, merged_map = clinical_forms.combine_sources(sources)
    draft["combined_text"] = combined_text
    draft["merged_phi_map"] = merged_map

    plans = retrieval_planner.plan(spec, combined_text)
    house_style: dict[str, list[str]] = {}
    for key, field_plan in plans.items():
        if not field_plan.want_exemplars:
            continue
        hits = exemplars.retrieve(spec.form_id, key, field_plan.query)
        if hits:
            house_style[key] = hits

    placeholder = st.empty()
    started = time.monotonic()
    try:
        with st.spinner("Generating on this computer — this can take a minute. Nothing leaves your device."):
            _, backend, _label = backends.select_backend()
            chunks = clinical_forms.generate_form_document(
                combined_text, spec, backend, stream=True, phi_values=phi_values,
                acknowledged=acknowledged, exemplars=house_style,
            )
            raw = _stream_into(placeholder, chunks, started)
    except (carenotes.CareNoteError, backends.BackendError) as exc:
        placeholder.empty()
        st.error(str(exc))
        return

    placeholder.empty()
    draft["deidentified"] = raw
    draft["field_values"] = clinical_forms.parse_fields(spec, raw)
    _invalidate_form_export(draft)
    draft["history"] = []
    st.rerun()


def render_form_draft(docs: dict, selected_names: list[str], spec, draft: dict) -> None:
    from carescribe.core import clinical_forms, exemplars

    st.markdown("#### Draft (de-identified)")
    st.caption("Still contains placeholders — safe to display, share, and save.")
    st.markdown(clinical_forms.render_preview(spec, draft["field_values"]))

    saved = exemplars.count(spec.form_id)
    cols = st.columns([3, 2])
    with cols[0]:
        if st.button("★ Save as house-style example", key=f"form_exemplar_{id(draft)}"):
            try:
                exemplars.add_exemplar(spec.form_id, draft["field_values"])
                st.success("Saved. Future drafts of this form will match its style.")
            except exemplars.ExemplarError as exc:
                st.warning(str(exc))
            st.rerun()
    with cols[1]:
        if saved:
            st.caption(f"{saved} house-style example{'s' if saved != 1 else ''} on file — "
                       "used to steer generation.")

    _render_reference_panel(spec, draft.get("combined_text", ""))

    render_form_refinement(docs, selected_names, spec, draft)
    render_form_reidentification(spec, draft)


def render_form_refinement(docs: dict, selected_names: list[str], spec, draft: dict) -> None:
    from carescribe.core import clinical_forms

    with st.expander("Refine this draft", expanded=False):
        st.caption("Refinement runs on the same de-identified source and the current draft.")
        for instruction, _ in draft["history"]:
            st.markdown(f"- _{instruction}_")
        instruction = st.text_input(
            "What would you like changed?", key=f"form_refine_{id(draft)}",
            placeholder="e.g. expand the risk assessment section",
        )
        phi_values = [v for name in selected_names for v in docs[name].phi_map.values()]
        acknowledged = [v for name in selected_names for v in docs[name].dismissed]
        status = ollama_client.status()
        if st.button(
            "Apply", key=f"form_refine_go_{id(draft)}",
            disabled=not instruction or not status["models"],
        ):
            placeholder = st.empty()
            started = time.monotonic()
            try:
                with st.spinner("Revising…"):
                    chunks = clinical_forms.refine_form_document(
                        draft["combined_text"], draft["deidentified"], instruction, spec,
                        backends.select_backend()[1], stream=True,
                        history=draft["history"], phi_values=phi_values,
                        acknowledged=acknowledged,
                    )
                    revised = _stream_into(placeholder, chunks, started)
            except (carenotes.CareNoteError, backends.BackendError) as exc:
                placeholder.empty()
                st.error(str(exc))
                return
            placeholder.empty()
            draft["history"].append((instruction, ""))
            draft["deidentified"] = revised
            draft["field_values"] = clinical_forms.parse_fields(spec, revised)
            _invalidate_form_export(draft)
            st.rerun()


def render_form_reidentification(spec, draft: dict) -> None:
    from carescribe.core import clinical_forms

    st.markdown("#### Re-identify and export (local only)")
    st.warning(
        "**This produces a document containing real patient identifiers.** "
        "It happens entirely in Python on this machine."
    )

    merged_map = draft.get("merged_phi_map", {})
    if st.button("🔓 Re-identify and fill the form", key=f"form_reid_{id(draft)}", disabled=not merged_map):
        resolved_fields = {}
        unresolved: list[str] = []
        for key, text in draft["field_values"].items():
            resolved_text, field_unresolved = carenotes.finalise(text, merged_map)
            resolved_fields[key] = resolved_text
            unresolved.extend(field_unresolved)
        draft["resolved_field_values"] = resolved_fields
        draft["unresolved"] = sorted(set(unresolved))
        st.rerun()

    if draft["unresolved"]:
        st.error(
            "**Blocked — these placeholders could not be resolved:** "
            + ", ".join(f"`{token}`" for token in draft["unresolved"])
        )
        return

    if draft.get("resolved_field_values"):
        st.success("Re-identified. Every placeholder resolved.")
        header_values = draft.get("header_values", {})
        output = clinical_forms.fill_template(spec, draft["resolved_field_values"], header_values)
        st.download_button(
            "⬇ .docx (contains PHI)", output,
            file_name=f"{batch.safe_stem(spec.title)}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"form_dl_{id(draft)}",
        )


def _as_docx(text: str) -> bytes:
    """Render generated text to a .docx in memory — nothing touches disk."""
    from docx import Document as DocxDocument

    buffer = io.BytesIO()
    doc = DocxDocument()
    for line in (text or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            doc.add_paragraph(line)
    doc.save(buffer)
    return buffer.getvalue()


def section_handoff() -> None:
    docs = documents()
    st.subheader("5. Generate report")

    mode = st.radio(
        "What do you want to generate?",
        ["Free-form note", "Clinical form"],
        horizontal=True, key="generation_mode",
    )
    if mode == "Clinical form":
        render_clinical_form_panel(docs)
        return

    approved = [doc for doc in docs.values() if doc.approved]
    if not approved:
        st.button(
            "Generate report", disabled=True,
            help=carenotes.DISABLED_MESSAGE, use_container_width=False,
        )
        st.info(carenotes.DISABLED_MESSAGE)
        st.caption(
            "Generation runs only on approved text. The model receives the "
            "de-identified document and nothing else — the identity mapping "
            "has no path out of this process."
        )
        if not generation_status.generation_status().ready:
            st.divider()
            render_setup_card()
        return

    st.caption(
        f"{len(approved)} approved document(s). The model receives the "
        "de-identified text only; re-identification happens afterwards, "
        "locally, in Python."
    )
    chosen = st.selectbox(
        "Document", [doc.name for doc in approved], key="gen_doc"
    )
    document = documents().get(chosen)
    if document is not None:
        render_generation_panel(document)


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------

def _pipeline_step(docs: dict) -> int:
    """The 0-based step the reviewer stands on, for components.step_tracker()."""
    if not docs:
        return 0
    values = list(docs.values())
    if not all(d.analyzed or d.error for d in values):
        return 1
    if not any(d.approved for d in values):
        return 2
    if not all(d.approved or d.error for d in values):
        return 3
    return 4


def _privacy_state() -> str:
    if backends.cloud_enabled():
        return "cloud"
    if st.session_state.get("downloading_model"):
        return "downloading"
    return "offline"


def main() -> None:
    ui_theme.inject()
    render_sidebar()

    st.markdown(
        ui.hero(
            "CareScribe — de-identification & review",
            "Load a batch, de-identify each document locally on the CPU, review "
            "and correct what was found, then approve. Nothing leaves this machine.",
            _privacy_state(),
        ),
        unsafe_allow_html=True,
    )

    # Load the model before anything can ask for it, so the first click is
    # never the thing that triggers a silent multi-second load.
    engine_state = ensure_engine_ready()
    if engine_state["error"]:
        render_engine_failure(engine_state)
        return

    docs = documents()
    order = st.session_state.order
    st.markdown(ui.step_tracker(_pipeline_step(docs)), unsafe_allow_html=True)

    # Each numbered step gets its own card (see ui.theme.CSS). Sections still
    # guard themselves — the predicates here only decide whether to open a card,
    # so a step that has nothing to show never leaves an empty box behind.
    has_review = bool(docs) and any(
        docs[name].analyzed and not docs[name].error for name in order
    )
    steps = [(section_load, True), (section_process, bool(docs)),
             (section_review, has_review), (section_batch_status, bool(docs)),
             (section_handoff, True)]
    for section, show in steps:
        if not show:
            continue
        with st.container(border=True):
            # Marks this border wrapper as a step card for ui.theme.CSS; hidden.
            st.markdown('<span class="cs-card"></span>', unsafe_allow_html=True)
            section()

    st.divider()
    st.caption(
        f"Model: `{engine_state['model']}` (loaded in "
        f"{engine_state['elapsed']:.1f}s) · Log file: `{applog.log_path()}`"
    )


def render_unexpected_error(exc: BaseException) -> None:
    """The last line of defence: a calm message instead of a stack trace.

    A clinician who sees a Python traceback in a medical tool has no idea
    whether their patient data is safe, and reasonably stops trusting the app.
    The traceback still goes to the log — where it is useful — and the screen
    gets a sentence and a next step.
    """
    applog.exception("unhandled exception in the UI")
    st.error(
        "**Something went wrong.**\n\n"
        "Your documents have not been sent anywhere and nothing has been "
        "changed on disk. You can usually carry on by going back a step, or "
        "restart CareScribe."
    )
    st.caption(
        "If it keeps happening, send this log file — it contains timings and "
        "file sizes, no patient information:"
    )
    st.code(str(applog.log_path()), language="text")
    with st.expander("Technical detail (for whoever supports this app)"):
        st.code(f"{type(exc).__name__}: {exc}", language="text")



applog.log("app start frozen=%s", desktop.is_frozen())

try:
    main()
except Exception as _exc:  # noqa: BLE001 — the boundary catches everything
    render_unexpected_error(_exc)
